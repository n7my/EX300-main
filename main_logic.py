import os
import sys
import threading
# import win32api
# import win32print
import CAN_option
# from main_ui import Ui_Form
from ctypes import *
import time
from queue import Queue
from PyQt5.QtCore import *
from PyQt5.QtGui import *
from PyQt5.QtWidgets import *
from CAN_option import *
from abc import ABCMeta, abstractmethod
# from workThread import WorkThread
from new1440_all import Ui_Form

import psutil
import xlwt
import signal
from PyQt5 import QtCore, QtGui, QtWidgets

import threading
import queue
import serial.tools.list_ports
import traceback
import gif_QMessageBox

q = queue.Queue()
event_canInit = threading.Event()
isMainRunning = True


class Ui_Control(QMainWindow, Ui_Form):
    #                -/+10V       0-10V       0-5V        1-5V       4-20mA      0-20mA      -/+5V
    AORangeArray = [0x0b, 0x00, 0xe8, 0x03, 0xe9, 0x03, 0xea, 0x03, 0x15, 0x00, 0xeb, 0x03, 0xec, 0x03]
    AIRangeArray = [0x29, 0x00, 0x2a, 0x00, 0x10, 0x27, 0x11, 0x27, 0x33, 0x00, 0x34, 0x00, 0x12, 0x27]

    currentTheory = [0x6c00, (int)(0x6c00 * 0.75), (int)(0x6c00 * 0.5), (int)(0x6c00 * 0.25), 0]  # 0x6c00 =>27648
    arrayCur = ["20mA测试", "15mA测试", "10mA测试", "5mA测试", "0mA测试"]
    highCurrent = 59849
    lowCurrent = 32768

    voltageTheory = [0x6c00, 0x3600, 0x00, 0xca00, 0x9400]  # 0x6c00 =>27648
    arrayVol = ["10V测试", "5V测试", "0V测试", "-5V测试", "-10V测试"]

    highVoltage = 59849
    lowVoltage = 5687

    HORIZONTAL_LINE = "\n------------------------------------------------------------------------------------------------------------\n\n"

    module_type = ''
    module_pn = ''
    module_sn = ''
    module_rev = ''
    # AI模块CAN地址
    CANAddr_AI = 1
    # AO模块CAN地址
    CANAddr_AO = 2
    # DI模块CAN地址
    CANAddr_DI = 1
    # DO模块CAN地址
    CANAddr_DO = 2
    # 继电器CAN地址
    CANAddr_relay = 3
    # 波特率
    baud_rate = 1000
    # 信号等待时间
    waiting_time = 5000
    # 标定接收信号次数
    receive_num = 8
    # 循环次数
    loop_num = 1
    pause_num = 1

    # 生成表格的相关标识
    # DI测试是否通过
    isDIPassTest = True
    # DO测试是否通过
    isDOPassTest = True
    # DO通道数据初始化
    DO_channelData = 0
    # DI通道数据初始化
    DI_channelData = 0
    # DO通道错误数据记录
    DODataCheck = [True for i in range(32)]
    # DI通道错误数据记录
    DIDataCheck = [True for i in range(32)]

    # 是否标定
    isCalibrate = False
    # 是否标定电压
    isCalibrateVol = False
    # 是否标定电流
    isCalibrateCur = False

    # 是否检测
    isCalibrate = False
    # 是否检测AI电压
    isAITestVol = False
    # 是否检测AI电流
    isAITestCur = False
    # 是否检测AO电压
    isAOTestVol = False
    # 是否检测AO电流
    isAOTestCur = False
    # 是否检测CAN_Run+CAN_Error
    isTestCANRunErr = False
    # 是否检测Run+Error
    isTestRunErr = False
    # AI测试是否通过
    isAIPassTest = True
    isAIVolPass = True
    isAICurPass = True
    # AO测试是否通过
    isAOPassTest = True
    isAOVolPass = True
    isAOCurPass = True

    isLEDRunOK = True
    isLEDErrOK = True
    isLEDPass = True

    # CAN帧结构体
    # 发送帧ID
    TRANSMIT_ID = 0x602

    # 接收帧ID
    RECEIVE_ID = 0x281

    # 时间标识
    TIME_STAMP = 0

    # 是否使用时间标识
    TIME_FLAG = 0

    # 发送帧类型
    TRANSMIT_SEND_TYPE = 1

    # 接收帧类型
    RECEIVE_SEND_TYPE = 1

    # 是否是远程帧
    REMOTE_FLAG = 0

    # 是否是扩展帧
    EXTERN_FLAG = 0

    # 数据长度DLC
    DATA_LEN = 8

    # 用来接收的帧结构体数组的长度, 适配器中为每个通道设置了2000帧左右的接收缓存区
    RECEIVE_LEN = 1

    # 接收保留字段
    WAIT_TIME = 500

    # 要发送的帧结构体数组的长度(发送的帧数量), 最大为1000, 建议设为1, 每次发送单帧, 以提高发送效率
    TRANSMIT_LEN = 1

    ubyte_array_8 = c_ubyte * 8
    DATA = ubyte_array_8(0, 0, 0, 0, 0, 0, 0, 0)
    ubyte_array_3 = c_ubyte * 3
    RESERVED_3 = ubyte_array_3(0, 0, 0)
    m_can_obj = CAN_option.VCI_CAN_OBJ(RECEIVE_ID, TIME_STAMP, TIME_FLAG, RECEIVE_SEND_TYPE,
                                       REMOTE_FLAG, EXTERN_FLAG, DATA_LEN, DATA, RESERVED_3)

    # 测试总数
    testNum = 0
    # 测试类型
    testType = {}
    # 模块信息
    inf = {'AO2': '待检AO模块', 'AI1': '配套AI模块', 'AI2': '待检AI模块', 'AO1': '配套AO模块',
           'DO2': '待检DO模块', 'DI1': '配套DI模块', 'DI2': '待检DI模块', 'DO1': '配套DO模块'}
    module_1 = ''
    module_2 = ''
    # 通道数
    m_Channels = 0
    AI_Channels = 0
    AO_Channels = 0
    # 发送的数据
    ubyte_array_transmit = c_ubyte * 8
    m_transmitData = ubyte_array_transmit(0x00, 0x00, 0x00, 0x00, 0x00, 0x00, 0x00, 0x00)
    # 接收的数据
    ubyte_array_receive = c_ubyte * 8
    m_receiveData = ubyte_array_receive(0x00, 0x00, 0x00, 0x00, 0x00, 0x00, 0x00, 0x00)
    CAN_errorLED = True
    CAN_runLED = True
    errorLED = True
    runLED = True
    errorNum = 0
    errorInf = ''
    isAllScreen = False
    volReceValue = [0, 0, 0, 0, 0]
    curReceValue = [0, 0, 0, 0, 0]
    volPrecision = [0, 0, 0, 0, 0]
    curPrecision = [0, 0, 0, 0, 0]
    chPrecision = [0, 0, 0, 0]

    subRunFlag = True
    # 创建队列用于主线程和子线程之间的通信
    result_queue = Queue()

    # 停止CAN接收信号
    stop_signal = pyqtSignal(bool)

    # 3码转换后的ASCII码
    asciiCode_pn = []
    asciiCode_sn = []
    asciiCode_rev = []
    testFlag = ''

    config_param = {}

    current_dir = os.getcwd().replace('\\', '/') + "/_internal"

    # current_dir = os.getcwd().replace('\\', '/')
    def __init__(self, parent=None):
        super(Ui_Control, self).__init__(parent)
        self.setupUi(self)
        self.pushButton_13.setVisible(False)
        # self.label_11.setPixmap(QtGui.QPixmap(f"{current_dir}/beast5.png"))
        # self.label_11.setVisible(False)
        # CPU页面参数配置
        self.CPU_lineEdit_array = [self.lineEdit_33, self.lineEdit_34, self.lineEdit_35, self.lineEdit_36,
                                   self.lineEdit_37, self.lineEdit_38]
        self.CPU_lineEditName_array = ["CPU_IP", "工装1", "工装2", "工装3", "工装4", "工装5"]
        self.CPU_checkBox_array = [self.checkBox_50, self.checkBox_51, self.checkBox_52, self.checkBox_53,
                                   self.checkBox_55, self.checkBox_56, self.checkBox_57,
                                   self.checkBox_58, self.checkBox_59, self.checkBox_60, self.checkBox_61,
                                   self.checkBox_62, self.checkBox_63, self.checkBox_64, self.checkBox_65,
                                   self.checkBox_66, self.checkBox_54, self.checkBox_67, self.checkBox_68,
                                   self.checkBox_71, self.checkBox_73]
        self.CPU_checkBoxName_array = ["外观检测", "型号检查", "SRAM", "FLASH", "拨杆测试",
                                       "MFK按键", "掉电保存", "RTC测试", "FPGA",
                                       "各指示灯", "本体IN", "本体OUT", "以太网", "RS-232C",
                                       "RS-485",
                                       "右扩CAN", "MAC/三码写入", "MA0202", "测试报告", "修改参数", "全选"]

        self.setWindowFlags(Qt.FramelessWindowHint)  # 无边框窗口
        self.label_6.mousePressEvent = self.label_mousePressEvent
        self.label_6.mouseMoveEvent = self.label_mouseMoveEvent
        self.label.setStyleSheet(self.testState_qss['stop'])
        self.label.setText('UNTESTED')
        self.label_41.setAlignment(QtCore.Qt.AlignLeft | Qt.AlignVCenter)
        # 屏蔽pushbutton_9
        self.pushButton_9.setVisible(False)
        self.getSerialInf(0, type='CPU')
        # 读页面配置
        # self.loadConfig()必须放在类似comboBox.currentIndexChanged.connect(self.saveConfig)的代码之前
        # 某则一修改参数就会触发saveConfig，导致还没修改的参数被默认参数覆盖。
        self.loadConfig()
        if self.checkBox_73.isChecked():
            self.checkBox_73.setText("取消全选")
        else:
            self.checkBox_73.setText("全选")

        # 显示当前时间
        self.update_time()
        # self.lineEdit_PN = 'PRDr9HPA06Mz-00'
        # self.lineEdit_SN = 'S1223-001083'
        # self.lineEdit_REV = '06'
        self.tabWidget.setTabPosition(0)
        self.tabWidget.setGeometry(QtCore.QRect(750, 314, 672, 530))
        self.tab_bar = self.tabWidget.tabBar()
        # self.textBrowser_5.setStyleSheet("""
        #                     QTextBrowser::pane {
        #                         border: 1px solid #11826C;
        #                         padding: 2px;
        #                     }
        #                 """)
        # for tW in [self.tableWidget_AI,self.tableWidget_AO,self.tableWidget_DIDO]:
        #     tW.setStyleSheet("""
        #                         QtWidgets::pane {
        #                             border: 1px solid #11826C;
        #                             padding: 2px;
        #                         }
        #                     """)

        for tW in [self.tableWidget_AI, self.tableWidget_AO, self.tableWidget_DIDO, self.tableWidget_CPU]:
            tW.setEditTriggers(QAbstractItemView.NoEditTriggers)

        # 批量设置lineEdit只读
        lE_arr = [self.lineEdit, self.lineEdit_3, self.lineEdit_5, self.lineEdit_20, self.lineEdit_21, self.lineEdit_22,
                  self.lineEdit_45, self.lineEdit_46, self.lineEdit_47, self.lineEdit_30, self.lineEdit_31,
                  self.lineEdit_32
            , self.lineEdit_33, self.lineEdit_MA0202_SN, self.lineEdit_MA0202_PN, self.lineEdit_MA0202_REV]
        for lE in lE_arr:
            lE.setReadOnly(True)

        # 设置选中和未选中状态下的样式
        self.selected_color = QColor("#11826C")
        self.unselected_color = QColor("white")
        self.tabWidget.setStyleSheet("""
                    QTabWidget::pane {
                        border: 1px solid #11826C;
                        padding: 2px;
                    }
                """)
        self.tab_bar.setStyleSheet(f"""
                    QTabBar::tab:selected {{
                        background-color: #F5BD3D;
                        color: #212B36;
                        border-top-left-radius: 0px ;
                        border-top-right-radius: 0px;
                        padding: 8px;
                        width: 105px; 
                        height: 17px;

                    }}
                    QTabBar::tab:!selected {{
                        background-color: #11826C;
                        color: white;
                        border-top-left-radius: 0px;
                        border-top-right-radius: 0px;
                        padding: 8px;
                        width: 105px; 
                        height: 13px;
                        margin-top: 5px
                    }}
                    QTabBar::tab:first {{
                        margin-left: 0;
                    }}
                    QTabBar::tab:last {{
                        margin-right: 0;
                    }}
                    QTabBar::tab {{
                        padding: 5px;
                        margin-right: 2px;
                        margin-bottom: 0px;
                    }}
                    QTabBar::tab:hover {{
                        background-color: #F5BD3D;
                        color: #212B36;
                        border-top-left-radius: 0px ;
                        border-top-right-radius: 0px;
                        padding: 8px;
                        width: 105px; 
                        height: 17px;
                    }}
                 """)

        # 设置字体为黑体加粗
        self.tabBarFont = QFont()
        self.tabBarFont.setFamily("Arial")
        self.tabBarFont.setBold(True)
        self.tabBarFont.setPointSize(12)
        self.tab_bar.setFont(self.tabBarFont)

        # 设置按钮悬停时的样式
        self.hover_color = QColor("#11826C").lighter(150)
        listPushButton = [self.pushButton_5, self.pushButton_6, self.pushButton_10, self.pushButton_CPU_renewSerial]
        for pB in listPushButton:
            pB.setStyleSheet(f"""
                        QPushButton {{
                            background-color: #11826C;
                            color:white
                        }}
                        QPushButton:hover {{
                            background-color: {self.hover_color.name()};
                            color:#212B36
                        }}
                    """)

        self.dragging = False
        self.offset = QPoint()

        self.lineEdit_PN.setReadOnly(True)
        self.pushButton_8.clicked.connect(self.killUi)
        self.pushButton_11.clicked.connect(self.showMinimized)
        self.pushButton_11.clicked.connect(self.saveConfig)
        self.tableWidget_array = [self.tableWidget_DIDO, self.tableWidget_AI,
                                  self.tableWidget_AO, self.tableWidget_CPU]
        for tW in self.tableWidget_array:
            tW.setStyleSheet("background-color: rgb(255, 255, 255);")
        # self.tableWidget_AI.setStyleSheet("background-color: rgb(255, 255, 255);")
        # self.tableWidget_AO.setStyleSheet("background-color: rgb(255, 255, 255);")
        # self.tableWidget_CPU.setStyleSheet("background-color: rgb(255, 255, 255);")
        self.pushButton_4.clicked.connect(self.start_button)
        self.pushButton_3.clicked.connect(self.stop_button)

        self.pushButton_pause.clicked.connect(self.pause_button)

        self.pushButton_resume.clicked.connect(self.resume_button)

        self.pushButton.setEnabled(True)
        self.pushButton.setStyleSheet(self.topButton_qss['on'])

        self.pushButton_2.setEnabled(True)
        self.pushButton_2.setStyleSheet(self.topButton_qss['on'])

        listTopButton = [self.pushButton, self.pushButton_2]
        for tB in listTopButton:
            tB.setStyleSheet(f"""
                                    QPushButton {{
                                        background-color: #11826C;
                                        color:white;
                                        font: bold 25pt \"宋体\"
                                    }}
                                    QPushButton:hover {{
                                        background-color: {self.hover_color.name()};
                                        color:#212B36;
                                        font: bold 25pt \"宋体\"
                                    }}
                                """)

        self.pushButton_3.setEnabled(False)
        self.pushButton_3.setStyleSheet(self.topButton_qss['off'])

        self.pushButton_4.setEnabled(False)
        self.pushButton_4.setStyleSheet(self.topButton_qss['off'])
        # # 设置按钮悬停时的样式
        # self.hover_pause = QColor("#ffff00").lighter(50)
        self.pushButton_pause.setStyleSheet(self.topButton_qss['pause_on'])
        self.pushButton_pause.setStyleSheet(f"""
                                QPushButton {{
                                    {self.topButton_qss['pause_on']}
                                }}
                                QPushButton:hover {{
                                    background-color: rgb(255, 255, 0);
                                    color:rgb(0, 0, 0)
                                }}
                            """)
        self.pushButton_pause.setEnabled(False)
        self.pushButton_pause.setVisible(False)
        # 设置按钮悬停时的样式
        self.hover_resume = QColor("#00ff00").lighter(50)
        self.pushButton_resume.setStyleSheet(self.topButton_qss['start_on'])
        self.pushButton_resume.setStyleSheet(f"""
                        QPushButton {{
                            {self.topButton_qss['start_on']}
                        }}
                        QPushButton:hover {{
                            background-color: rgb(0, 255, 0);
                            color:rgb(0, 0, 0)
                        }}
                    """)
        self.pushButton_resume.setEnabled(False)
        self.pushButton_resume.setVisible(False)
        for tW in self.tableWidget_array:
            tW.horizontalHeader().setStyleSheet("background-color: #11826C;")
        # self.tableWidget_AO.horizontalHeader().setStyleSheet("background-color: #11826C;")
        # self.tableWidget_DIDO.horizontalHeader().setStyleSheet("background-color: #11826C;")
        # self.tableWidget_CPU.horizontalHeader().setStyleSheet("background-color: #11826C;")

        self.checkBox_5.setEnabled(not self.radioButton.isChecked())
        self.checkBox_6.setEnabled(not self.radioButton.isChecked())
        self.checkBox_7.setEnabled(self.radioButton_5.isChecked())
        self.checkBox_8.setEnabled(self.radioButton_5.isChecked())
        self.checkBox_9.setEnabled(self.radioButton_5.isChecked())
        self.checkBox_10.setEnabled(self.radioButton_5.isChecked())

        self.checkBox_29.setEnabled(self.radioButton_17.isChecked())
        self.checkBox_30.setEnabled(self.radioButton_17.isChecked())
        self.checkBox_31.setEnabled(self.radioButton_17.isChecked())
        self.checkBox_32.setEnabled(self.radioButton_17.isChecked())

        self.checkBox_35.setEnabled(not self.radioButton_18.isChecked())
        self.checkBox_36.setEnabled(not self.radioButton_18.isChecked())

        self.pushButton_5.clicked.connect(self.changeSaveDir)
        self.pushButton_6.clicked.connect(self.openSaveDir)

        self.tabIndex = self.tabWidget.currentIndex()
        if self.tabIndex == 0:
            self.tableWidget_DIDO.setVisible(True)
            self.tableWidget_AI.setVisible(False)
            self.tableWidget_AO.setVisible(False)
            self.tableWidget_CPU.setVisible(False)
            self.label_7.setText("DI/DO")
            if self.comboBox.currentIndex() == 1:
                self.lineEdit_PN.setText('P0000010391877')
            elif self.comboBox.currentIndex() == 4:
                self.lineEdit_PN.setText('P0000010392086')
            elif self.comboBox.currentIndex() == 5:
                self.lineEdit_PN.setText('P0000010392121')
            self.lineEdit.setText(self.lineEdit_PN.text())
        elif self.tabIndex == 1:
            self.tableWidget_DIDO.setVisible(False)
            self.tableWidget_AI.setVisible(True)
            self.tableWidget_AO.setVisible(False)
            self.tableWidget_CPU.setVisible(False)
            self.label_7.setText("AI")
            if self.comboBox_3.currentIndex() == 0:
                self.lineEdit_PN.setText('P0000010392361')
            self.lineEdit_22.setText(self.lineEdit_PN.text())
        elif self.tabIndex == 2:
            self.tableWidget_DIDO.setVisible(False)
            self.tableWidget_AI.setVisible(False)
            self.tableWidget_AO.setVisible(True)
            self.tableWidget_CPU.setVisible(False)
            self.label_7.setText("AO")
            if self.comboBox_5.currentIndex() == 0:
                self.lineEdit_PN.setText('P0000010392365')
            self.lineEdit_47.setText(self.lineEdit_PN.text())
        elif self.tabIndex == 3:
            self.tableWidget_DIDO.setVisible(False)
            self.tableWidget_AI.setVisible(False)
            self.tableWidget_AO.setVisible(False)
            self.tableWidget_CPU.setVisible(True)
            self.label_7.setText("CPU")
            if self.comboBox_20.currentIndex() == 0:
                self.lineEdit_PN.setText('P0000010390631')
            self.lineEdit_30.setText(self.lineEdit_PN.text())
        elif self.tabIndex == 4:
            self.tableWidget_DIDO.setVisible(False)
            self.tableWidget_AI.setVisible(False)
            self.tableWidget_AO.setVisible(False)
            self.tableWidget_CPU.setVisible(True)
            self.label_7.setText("CPU")
            # if self.comboBox_20.currentIndex()==0:
            self.lineEdit_PN.setText('MA0202')
            self.lineEdit_MA0202_PN.setText(self.lineEdit_PN.text())
        self.saveDir = self.label_41.text()
        self.tabWidget.currentChanged.connect(self.tabChange)

        self.pushButton_9.setEnabled(True)
        self.pushButton_9.clicked.connect(self.uiResize)

        # self.checkBox_27.setChecked(False)
        self.checkBox_27.stateChanged.connect(self.DIDOCANAddr_stateChanged)
        self.lineEdit_6.setEnabled(self.checkBox_27.isChecked())
        self.label_16.setEnabled(self.checkBox_27.isChecked())
        self.lineEdit_7.setEnabled(self.checkBox_27.isChecked())
        self.label_17.setEnabled(self.checkBox_27.isChecked())

        # self.checkBox_28.setChecked(False)
        self.checkBox_28.stateChanged.connect(self.DIDOAdPara_stateChanged)
        self.lineEdit_13.setEnabled(self.checkBox_28.isChecked())
        self.label_22.setEnabled(self.checkBox_28.isChecked())
        self.lineEdit_12.setEnabled(self.checkBox_28.isChecked())
        self.label_23.setEnabled(self.checkBox_28.isChecked())
        self.lineEdit_14.setEnabled(self.checkBox_28.isChecked())
        self.label_26.setEnabled(self.checkBox_28.isChecked())
        self.label_25.setEnabled(self.checkBox_28.isChecked())
        self.label_24.setEnabled(self.checkBox_28.isChecked())

        # self.checkBox_3.setChecked(False)
        self.checkBox_3.stateChanged.connect(self.AIAdPara_stateChanged)
        self.lineEdit_15.setEnabled(self.checkBox_3.isChecked())
        self.label_42.setEnabled(self.checkBox_3.isChecked())
        self.lineEdit_16.setEnabled(self.checkBox_3.isChecked())
        self.label_43.setEnabled(self.checkBox_3.isChecked())
        self.lineEdit_17.setEnabled(self.checkBox_3.isChecked())
        self.label_44.setEnabled(self.checkBox_3.isChecked())
        self.label_45.setEnabled(self.checkBox_3.isChecked())
        self.label_46.setEnabled(self.checkBox_3.isChecked())

        # self.checkBox_4.setChecked(False)
        self.checkBox_4.stateChanged.connect(self.AICANAddr_stateChanged)
        self.lineEdit_18.setEnabled(self.checkBox_4.isChecked())
        self.label_48.setEnabled(self.checkBox_4.isChecked())
        self.lineEdit_19.setEnabled(self.checkBox_4.isChecked())
        self.label_49.setEnabled(self.checkBox_4.isChecked())
        self.lineEdit_23.setEnabled(self.checkBox_4.isChecked())
        self.label_54.setEnabled(self.checkBox_4.isChecked())

        # self.radioButton.setChecked(True)
        self.radioButton.toggled.connect(self.AI_NotCalibrate)

        # self.radioButton_2.setChecked(False)
        self.radioButton_2.toggled.connect(self.AI_Calibrate)

        # self.radioButton_3.setChecked(False)
        self.radioButton_3.toggled.connect(self.AI_Calibrate)

        # self.radioButton_4.setChecked(True)
        self.radioButton_4.toggled.connect(self.AI_NotTest)

        # self.radioButton_5.setChecked(False)
        self.radioButton_5.toggled.connect(self.AI_Test)

        # self.radioButton_16.setChecked(True)
        self.radioButton_16.toggled.connect(self.AO_NotTest)

        self.horizontalLayout_52.addWidget(self.radioButton_17)
        self.radioButton_17.toggled.connect(self.AO_Test)

        # self.checkBox_33.setChecked(False)
        self.checkBox_33.stateChanged.connect(self.AOCANAddr_stateChanged)
        self.lineEdit_39.setEnabled(self.checkBox_33.isChecked())
        self.label_74.setEnabled(self.checkBox_33.isChecked())
        self.lineEdit_40.setEnabled(self.checkBox_33.isChecked())
        self.label_75.setEnabled(self.checkBox_33.isChecked())
        self.lineEdit_41.setEnabled(self.checkBox_33.isChecked())
        self.label_76.setEnabled(self.checkBox_33.isChecked())

        # self.checkBox_34.setChecked(False)
        self.checkBox_34.stateChanged.connect(self.AOAdPara_stateChanged)
        self.lineEdit_42.setEnabled(self.checkBox_34.isChecked())
        self.label_77.setEnabled(self.checkBox_34.isChecked())
        self.lineEdit_43.setEnabled(self.checkBox_34.isChecked())
        self.label_78.setEnabled(self.checkBox_34.isChecked())
        self.lineEdit_44.setEnabled(self.checkBox_34.isChecked())
        self.label_79.setEnabled(self.checkBox_34.isChecked())
        self.label_80.setEnabled(self.checkBox_34.isChecked())
        self.label_81.setEnabled(self.checkBox_34.isChecked())

        self.radioButton_18.toggled.connect(self.AO_NotCalibrate)

        # self.radioButton_19.setChecked(True)
        self.radioButton_19.toggled.connect(self.AO_Calibrate)

        self.radioButton_20.toggled.connect(self.AO_Calibrate)

        # self.pushButton_9.clicked.connect(self.sendMessage)
        self.pushbutton_allScreen.setStyleSheet(f"""
                                    QPushButton {{
                                        image: url({self.current_dir}/全屏按钮.png);

                                    }}
                                    QPushButton:hover {{
                                        image: url({self.current_dir}/全屏按钮2.png);
                                    }}
                                """)
        self.pushbutton_cancelAllScreen.setStyleSheet(f"""
                                     QPushButton {{
                                         image: url({self.current_dir}/取消全屏按钮.png);

                                     }}
                                     QPushButton:hover {{
                                         image: url({self.current_dir}/取消全屏按钮2.png);
                                     }}
                                 """)
        self.pushButton_8.setStyleSheet(f"""
                                     QPushButton {{
                                         font: 75 26pt \"Arial\";
                                         background-color:  #11826C;
                                         color: rgb(255, 255, 255);
                                     }}
                                     QPushButton:hover {{
                                         font: 75 26pt \"Arial\";
                                         background-color:  #F5BD3D;
                                         color: rgb(0, 0, 0);

                                     }}
                                 """)
        self.pushButton_11.setStyleSheet(f"""
                                             QPushButton {{
                                                 font: 75 26pt \"Arial\";
                                                 background-color:  #11826C;
                                                 color: rgb(255, 255, 255);
                                             }}
                                             QPushButton:hover {{
                                                 font: 75 26pt \"Arial\";
                                                 background-color:  #F5BD3D;
                                                 color: rgb(0, 0, 0);

                                             }}
                                         """)

        self.pushbutton_allScreen.clicked.connect(self.allScreen)

        self.pushbutton_cancelAllScreen.setEnabled(False)
        self.pushbutton_cancelAllScreen.setVisible(False)
        self.pushbutton_cancelAllScreen.clicked.connect(self.cancelAllScreen)

        # CPU页面初始化
        self.CPU_comboBox_array = [self.comboBox_20, self.comboBox_21, self.comboBox_22, self.comboBox_23,
                                   self.comboBox_24, self.comboBox_25, self.comboBox_power]
        for comboBox in self.CPU_comboBox_array:
            comboBox.currentIndexChanged.connect(self.saveConfig)
        for lineEdit in self.CPU_lineEdit_array:
            lineEdit.textChanged.connect(self.saveConfig)
        for checkBox in self.CPU_checkBox_array:
            checkBox.toggled.connect(self.saveConfig)
        self.CPU_paramChanged()
        self.checkBox_71.stateChanged.connect(self.CPU_paramChanged)
        self.checkBox_73.stateChanged.connect(self.testAllorNot)
        self.lineEdit_33.setText('11-22-33-44-55-66')

        # self.lineEdit_PN.setPlaceholderText('请输入PN码')
        self.lineEdit_SN.setPlaceholderText('请输入SN码')
        # self.lineEdit_SN.setReadOnly(True)
        self.lineEdit_REV.setPlaceholderText('请输入REV码')
        self.lineEdit_REV.setReadOnly(True)
        self.lineEdit_MAC.setPlaceholderText('请输入MAC地址')
        self.lineEdit_MAC.setReadOnly(False)
        # self.lineEdit_PN.setFocus()
        # self.lineEdit_PN.editingFinished.connect(self.inputPN)
        self.lineEdit_SN.setFocus()
        self.lineEdit_SN.editingFinished.connect(self.inputSN)
        self.lineEdit_REV.editingFinished.connect(self.inputREV)
        self.pushButton_10.clicked.connect(self.reInputPNSNREV)

        self.lineEdit_PN.textChanged.connect(self.changeTabWidgetPN)

        self.radioButton_16.toggled.connect(self.setFocusLineEdit)
        self.radioButton_17.toggled.connect(self.setFocusLineEdit)
        self.radioButton_18.toggled.connect(self.setFocusLineEdit)
        self.radioButton_19.toggled.connect(self.setFocusLineEdit)
        self.radioButton_20.toggled.connect(self.setFocusLineEdit)
        self.checkBox_29.toggled.connect(self.setFocusLineEdit)
        self.checkBox_30.toggled.connect(self.setFocusLineEdit)
        self.checkBox_31.toggled.connect(self.setFocusLineEdit)
        self.checkBox_32.toggled.connect(self.setFocusLineEdit)
        self.checkBox_36.toggled.connect(self.setFocusLineEdit)
        self.checkBox_35.toggled.connect(self.setFocusLineEdit)

        self.radioButton.toggled.connect(self.saveConfig)
        self.radioButton_2.toggled.connect(self.saveConfig)
        self.radioButton_3.toggled.connect(self.saveConfig)
        self.radioButton_4.toggled.connect(self.saveConfig)
        self.radioButton_5.toggled.connect(self.saveConfig)
        self.radioButton_16.toggled.connect(self.saveConfig)
        self.radioButton_17.toggled.connect(self.saveConfig)
        self.radioButton_18.toggled.connect(self.saveConfig)
        self.radioButton_19.toggled.connect(self.saveConfig)
        self.radioButton_20.toggled.connect(self.saveConfig)
        self.checkBox_29.toggled.connect(self.saveConfig)
        self.checkBox_30.toggled.connect(self.saveConfig)
        self.checkBox_31.toggled.connect(self.saveConfig)
        self.checkBox_32.toggled.connect(self.saveConfig)
        self.checkBox_36.toggled.connect(self.saveConfig)
        self.checkBox_35.toggled.connect(self.saveConfig)
        self.checkBox.toggled.connect(self.saveConfig)
        self.checkBox_2.toggled.connect(self.saveConfig)
        self.checkBox_5.toggled.connect(self.saveConfig)
        self.checkBox_6.toggled.connect(self.saveConfig)
        self.checkBox_7.toggled.connect(self.saveConfig)
        self.checkBox_8.toggled.connect(self.saveConfig)
        self.checkBox_9.toggled.connect(self.saveConfig)
        self.checkBox_10.toggled.connect(self.saveConfig)
        self.comboBox.currentIndexChanged.connect(self.saveConfig)
        self.comboBox_3.currentIndexChanged.connect(self.saveConfig)
        self.comboBox_5.currentIndexChanged.connect(self.saveConfig)
        self.comboBox.currentIndexChanged.connect(self.changePN)
        self.comboBox_3.currentIndexChanged.connect(self.changePN)
        self.comboBox_5.currentIndexChanged.connect(self.changePN)

        self.lineEdit_6.textChanged.connect(self.saveConfig)
        self.lineEdit_7.textChanged.connect(self.saveConfig)
        self.lineEdit_12.textChanged.connect(self.saveConfig)
        self.lineEdit_13.textChanged.connect(self.saveConfig)
        self.lineEdit_14.textChanged.connect(self.saveConfig)
        self.lineEdit_15.textChanged.connect(self.saveConfig)
        self.lineEdit_16.textChanged.connect(self.saveConfig)
        self.lineEdit_17.textChanged.connect(self.saveConfig)
        self.lineEdit_18.textChanged.connect(self.saveConfig)
        self.lineEdit_19.textChanged.connect(self.saveConfig)
        self.lineEdit_23.textChanged.connect(self.saveConfig)
        self.lineEdit_39.textChanged.connect(self.saveConfig)
        self.lineEdit_40.textChanged.connect(self.saveConfig)
        self.lineEdit_41.textChanged.connect(self.saveConfig)
        self.lineEdit_42.textChanged.connect(self.saveConfig)
        self.lineEdit_43.textChanged.connect(self.saveConfig)
        self.lineEdit_44.textChanged.connect(self.saveConfig)
        self.tabWidget.currentChanged.connect(self.saveConfig)
        self.pushButton.clicked.connect(self.saveConfig)
        self.pushButton.clicked.connect(self.generateLabel_userName)
        self.pushButton.clicked.connect(self.generateCheckBox_admin)
        self.pushButton.clicked.connect(self.userLogin)

        self.pushButton_2.clicked.connect(self.saveConfig)
        self.pushButton_3.clicked.connect(self.saveConfig)
        self.pushButton_4.clicked.connect(self.saveConfig)

        self.pushButton_12.setEnabled(False)
        self.pushButton_12.setVisible(False)
        self.pushButton_12.clicked.connect(self.uiRecovery)
        self.pushButton_13.clicked.connect(self.option_pushButton13)

        # MA0202界面初始化
        self.radioButton_MA0202.setChecked(True)
        self.MA0202_paramChanged()
        self.checkBox_MA0202_para.stateChanged.connect(self.MA0202_paramChanged)

        # 更新串口
        self.pushButton_CPU_renewSerial.clicked.connect(lambda: self.getSerialInf(1, type='CPU'))
        self.pushButton_MA0202_renewSerial.clicked.connect(lambda: self.getSerialInf(1, type='MA0202'))
        global isMainRunning  # 程序运行标志
        isMainRunning = True

        # 可编程电源参数
        self.power_off = [0x01, 0x06, 0x00, 0x01, 0x00, 0x00, 0xD8, 0x0A]
        self.power_on = [0x01, 0x06, 0x00, 0x01, 0x00, 0x01, 0x19, 0xCA]
        self.vol_24v = [0x01, 0x10, 0x00, 0x20, 0x00, 0x02, 0x04, 0x00, 0x00, 0x5D, 0xC0, 0xC9, 0x77]
        self.cur_2a = [0x01, 0x10, 0x00, 0x22, 0x00, 0x02, 0x04, 0x00, 0x00, 0x4E, 0x20, 0x44, 0x16]
        self.appearance = True
        #CPU外观无需检测
        self.checkBox_50.setEnabled(False)

        #启动时提示员工登录信息
        self.workerChange(['start','测试人员登录', '请输入员工编号：'])
        self.group_list = [self.groupBox,self.groupBox_3,self.groupBox_4,self.groupBox_5,
                           self.groupBox_7,self.groupBox_10,self.groupBox_11,self.groupBox_12,
                           self.groupBox_13,
                           self.groupBox_8,self.groupBox_14,self.groupBox_15,self.groupBox_16,
                           self.groupBox_17,
                           self.groupBox_18,self.groupBox_19,
                           self.groupBox_20]
        self.adminState = False
        # 启动时默认不允许操作
        for g in self.group_list:
            g.setEnabled(False)

    def generateLabel_userName(self):
        # 员工登录账号
        self.label_userName = QLabel(f'<h2>123456<h2>')

    def generateCheckBox_admin(self):
        # 管理员权限
        self.checkBox_admin = QCheckBox('管理员权限')

    def generateCheckBox_messageBox_widget(self,list):
        self.checkBox_messageBox_widget = QWidget()
        self.checkBoxs = [0,0,0,0,
                          0,0,0,0,
                          0,0,0,0,
                          0,0,0,0]
        self.checkBox_status = [True, True, True, True,
                                True, True, True, True,
                                True, True, True, True,
                                True, True, True, True]

        # 创建四个QHBoxLayout布局管理器
        layout1 = QHBoxLayout()
        layout2 = QHBoxLayout()
        layout3 = QHBoxLayout()
        layout4 = QHBoxLayout()

        # 添加16个checkBox，分成4组，每组4个，左右分布
        for i in range(0, 7, 2):
            button = QCheckBox('0' + str(i))
            button.setObjectName('通道' + str(i))  # 通道0，2，4，6
            if list[2] == 'odd':
                button.setEnabled(False)
            self.checkBoxs[i]=button
            # self.checkBox_messageBox_widget.append(button)
            layout1.addWidget(button)
        for i in range(1, 8, 2):
            button = QCheckBox('0' + str(i))
            button.setObjectName('通道' + str(i))  # 通道1，3，5，7
            if list[2] == 'even':
                button.setEnabled(False)
            self.checkBoxs[i]=button
            # self.checkBox_messageBox_widget.append(button)
            layout2.addWidget(button)
        for i in range(10, 17, 2):
            button = QCheckBox(str(i))
            button.setObjectName('通道' + str(i))  # 通道10，12，14，16
            if list[2] == 'odd':
                button.setEnabled(False)
            self.checkBoxs[i-2]=button
            # self.checkBox_messageBox_widget.append(button)
            layout3.addWidget(button)
        for i in range(11, 18, 2):
            button = QCheckBox(str(i))
            button.setObjectName('通道' + str(i))  # 通道11，13，15，17
            if list[2] == 'even':
                button.setEnabled(False)
            self.checkBoxs[i-2]=button
            # self.checkBox_messageBox_widget.append(button)
            layout4.addWidget(button)

        # 创建一个QWidget对象，并将四个QHBoxLayout布局管理器添加到该对象中
        # widget = QWidget()
        layout = QVBoxLayout()

        layout.addLayout(layout1)
        layout.addLayout(layout2)
        layout.addLayout(layout3)
        layout.addLayout(layout4)

        self.checkBox_messageBox_widget.setLayout(layout)

        for element in self.checkBoxs:
            element.toggled.connect(self.update_checkBox_status)

    def update_checkBox_status(self):
        for i in range(len(self.checkBoxs)):
            if self.checkBoxs[i].isChecked():
                self.checkBox_status[i] = False
            else:
                self.checkBox_status[i] = True

    def generateDefaultConfigFile(self):
        config_str = "{'savePath': 'D:/MyData/wujun89/Desktop/EX300_x64_python',\n" \
                     "'currentIndex': 0,\n" \
                     "'AO_型号': 0,\n" \
                     "'AO_CAN_修改参数': False,\n" \
                     "'AO_CAN_工装': '1',\n" \
                     "'AO_CAN_待检': '2',\n" \
                     "'AO_CAN_继电器': '3',\n" \
                     "'AO_附加_修改参数': False,\n" \
                     "'AO_附加_波特率': '1000',\n" \
                     "'AO_附加_等待时间': '5000',\n" \
                     "'AO_附加_接收次数': '8',\n" \
                     "'AO_不标定': False,\n" \
                     "'AO_仅标定1': False,\n" \
                     "'AO_标定all': False,\n" \
                     "'AO_标定电压': False,\n" \
                     "'AO_标定电流': False,\n" \
                     "'AO_不检测': False,\n" \
                     "'AO_检测': False,\n" \
                     "'AO_检测电压': False,\n" \
                     "'AO_检测电流': False,\n" \
                     "'AO_检测CAN': False,\n" \
                     "'AO_检测run': False,\n" \
                     "'AI_型号': 0,\n" \
                     "'AI_CAN_修改参数': False,\n" \
                     "'AI_CAN_工装': '1',\n" \
                     "'AI_CAN_待检': '2',\n" \
                     "'AI_CAN_继电器': '3',\n" \
                     "'AI_附加_修改参数': False,\n" \
                     "'AI_附加_波特率': '1000',\n" \
                     "'AI_附加_等待时间': '5000',\n" \
                     "'AI_附加_接收次数': '8',\n" \
                     "'AI_不标定': False,\n" \
                     "'AI_仅标定1': False,\n" \
                     "'AI_标定all': False,\n" \
                     "'AI_标定电压': False,\n" \
                     "'AI_标定电流': False,\n" \
                     "'AI_不检测': False,\n" \
                     "'AI_检测': False,\n" \
                     "'AI_检测电压': False,\n" \
                     "'AI_检测电流': False,\n" \
                     "'AI_检测CAN': False,\n" \
                     "'AI_检测run': False,\n" \
                     "'DIDO_型号': 0,\n" \
                     "'DIDO_CAN_修改参数': False,\n" \
                     "'DIDO_CAN_工装': '1',\n" \
                     "'DIDO_CAN_待检': '2',\n" \
                     "'DIDO_附加_修改参数': False,\n" \
                     "'DIDO_附加_波特率': '1000',\n" \
                     "'DIDO_附加_间隔时间': '5000',\n" \
                     "'DIDO_附加_循环次数': '1',\n" \
                     "'DIDO_检测CAN': False,\n" \
                     "'DIDO_检测run': False,\n" \
                     " 'CPU_型号': 0, \n" \
                     " 'CPU_IP': '11-22-33-44-55-66', \n" \
                     "'CPU_232COM': 5,\n" \
                     "'CPU_485COM': 4,\n" \
                     "'CPU_typecCOM': 3,\n" \
                     "'工装1': '1',\n" \
                     "'工装2': '2',\n" \
                     "'工装3': '3',\n" \
                     "'工装4': '4',\n" \
                     "'工装5': '5',\n" \
                     "'外观检测': False,\n" \
                     "'型号检查': False,\n" \
                     "'SRAM': False,\n" \
                     "'FLASH': False,\n" \
                     "'MAC/三码写入': False,\n" \
                     "'FPGA': False,\n" \
                     "'拨杆测试': False,\n" \
                     "'MFK按键': False,\n" \
                     "'RTC测试': False,\n" \
                     "'掉电保存': False,\n" \
                     "'各指示灯': True,\n" \
                     "'本体IN': False,\n" \
                     "'本体OUT': False,\n" \
                     "'以太网': False,\n" \
                     "'RS-232C': False,\n" \
                     "'RS-485': False,\n" \
                     "'右扩CAN': False,\n" \
                     "'MA0202': False,\n" \
                     "'测试报告': True,\n" \
                     "'固件烧录': False,\n" \
                     "'U盘读写': False,\n" \
                     "'修改参数': False,\n" \
                     "'全选': False,\n" \
                     "'选项板232':0,\n" \
                     "'选项板485':0,\n" \
                     "'可编程电源':0}"
        self.configFile = open(f'{self.current_dir}/config.txt', 'w', encoding='utf-8')
        self.configFile.write(config_str)
        self.configFile.close()

    def loadConfig(self):
        try:
            with open(f'{self.current_dir}/config.txt', 'r+', encoding='utf-8') as file:
                config_content = file.read()
        except:
            self.showInf("配置文件不存在，初始化界面！" + self.HORIZONTAL_LINE)
            self.generateDefaultConfigFile()
            with open(f'{self.current_dir}/config.txt', 'r+', encoding='utf-8') as file:
                config_content = file.read()

        self.config_param = eval(config_content)
        self.label_41.setText(self.config_param["savePath"])
        self.tabWidget.setCurrentIndex(self.config_param["currentIndex"])

        # AO页面配置
        self.comboBox_5.setCurrentIndex(self.config_param["AO_型号"])
        self.checkBox_33.setChecked(self.config_param["AO_CAN_修改参数"])
        self.lineEdit_39.setText(self.config_param["AO_CAN_工装"])
        self.lineEdit_40.setText(self.config_param["AO_CAN_待检"])
        self.lineEdit_41.setText(self.config_param["AO_CAN_继电器"])
        self.checkBox_34.setChecked(self.config_param["AO_附加_修改参数"])
        self.lineEdit_42.setText(self.config_param["AO_附加_波特率"])
        self.lineEdit_43.setText(self.config_param["AO_附加_等待时间"])
        self.lineEdit_44.setText(self.config_param["AO_附加_接收次数"])
        self.radioButton_18.setChecked(self.config_param["AO_不标定"])
        self.radioButton_19.setChecked(self.config_param["AO_仅标定1"])
        self.radioButton_20.setChecked(self.config_param["AO_标定all"])
        self.checkBox_35.setChecked(self.config_param["AO_标定电压"])
        self.checkBox_36.setChecked(self.config_param["AO_标定电流"])
        self.radioButton_16.setChecked(self.config_param["AO_不检测"])
        self.radioButton_17.setChecked(self.config_param["AO_检测"])
        self.checkBox_29.setChecked(self.config_param["AO_检测电压"])
        self.checkBox_30.setChecked(self.config_param["AO_检测电流"])
        self.checkBox_31.setChecked(self.config_param["AO_检测CAN"])
        self.checkBox_32.setChecked(self.config_param["AO_检测run"])
        # AI页面配置
        self.comboBox_3.setCurrentIndex(self.config_param["AI_型号"])
        self.checkBox_4.setChecked(self.config_param["AI_CAN_修改参数"])
        self.lineEdit_18.setText(self.config_param["AI_CAN_工装"])
        self.lineEdit_19.setText(self.config_param["AI_CAN_待检"])
        self.lineEdit_23.setText(self.config_param["AI_CAN_继电器"])
        self.checkBox_3.setChecked(self.config_param["AI_附加_修改参数"])
        self.lineEdit_16.setText(self.config_param["AI_附加_波特率"])
        self.lineEdit_17.setText(self.config_param["AI_附加_等待时间"])
        self.lineEdit_15.setText(self.config_param["AI_附加_接收次数"])
        self.radioButton.setChecked(self.config_param["AI_不标定"])
        self.radioButton_2.setChecked(self.config_param["AI_仅标定1"])
        self.radioButton_3.setChecked(self.config_param["AI_标定all"])
        self.checkBox_5.setChecked(self.config_param["AI_标定电压"])
        self.checkBox_6.setChecked(self.config_param["AI_标定电流"])
        self.radioButton_4.setChecked(self.config_param["AI_不检测"])
        self.radioButton_5.setChecked(self.config_param["AI_检测"])
        self.checkBox_9.setChecked(self.config_param["AI_检测电压"])
        self.checkBox_10.setChecked(self.config_param["AI_检测电流"])
        self.checkBox_8.setChecked(self.config_param["AI_检测CAN"])
        self.checkBox_7.setChecked(self.config_param["AI_检测run"])
        # DIDO页面配置
        self.comboBox.setCurrentIndex(self.config_param["DIDO_型号"])
        self.checkBox_27.setChecked(self.config_param["DIDO_CAN_修改参数"])
        self.lineEdit_6.setText(self.config_param["DIDO_CAN_工装"])
        self.lineEdit_7.setText(self.config_param["DIDO_CAN_待检"])
        self.checkBox_28.setChecked(self.config_param["DIDO_附加_修改参数"])
        self.lineEdit_13.setText(self.config_param["DIDO_附加_波特率"])
        self.lineEdit_12.setText(self.config_param["DIDO_附加_间隔时间"])
        self.lineEdit_14.setText(self.config_param["DIDO_附加_循环次数"])
        self.checkBox.setChecked(self.config_param["DIDO_检测CAN"])
        self.checkBox_2.setChecked(self.config_param["DIDO_检测run"])

        # CPU页面配置
        self.comboBox_20.setCurrentIndex(self.config_param["CPU_型号"])
        self.comboBox_21.setCurrentIndex(self.config_param["CPU_232COM"])
        self.comboBox_22.setCurrentIndex(self.config_param["CPU_485COM"])
        self.comboBox_23.setCurrentIndex(self.config_param["CPU_typecCOM"])
        self.comboBox_24.setCurrentIndex(self.config_param["选项板232"])
        self.comboBox_25.setCurrentIndex(self.config_param["选项板485"])
        self.comboBox_power.setCurrentIndex(self.config_param["可编程电源"])

        for i in range(len(self.CPU_lineEdit_array)):
            self.CPU_lineEdit_array[i].setText(self.config_param[self.CPU_lineEditName_array[i]])

        for i in range(len(self.CPU_checkBox_array)):
            self.CPU_checkBox_array[i].setChecked(self.config_param[self.CPU_checkBoxName_array[i]])

    def saveConfig(self):
        self.config_param["savePath"] = self.label_41.text()
        self.config_param["currentIndex"] = self.tabWidget.currentIndex()

        # AO页面配置
        self.config_param["AO_型号"] = self.comboBox_5.currentIndex()
        self.config_param["AO_CAN_修改参数"] = self.checkBox_33.isChecked()
        self.config_param["AO_CAN_工装"] = self.lineEdit_39.text()
        self.config_param["AO_CAN_待检"] = self.lineEdit_40.text()
        self.config_param["AO_CAN_继电器"] = self.lineEdit_41.text()
        self.config_param["AO_附加_修改参数"] = self.checkBox_34.isChecked()
        self.config_param["AO_附加_波特率"] = self.lineEdit_42.text()
        self.config_param["AO_附加_等待时间"] = self.lineEdit_43.text()
        self.config_param["AO_附加_接收次数"] = self.lineEdit_44.text()
        self.config_param["AO_不标定"] = self.radioButton_18.isChecked()
        self.config_param["AO_仅标定1"] = self.radioButton_19.isChecked()
        self.config_param["AO_标定all"] = self.radioButton_20.isChecked()
        self.config_param["AO_标定电压"] = self.checkBox_35.isChecked()
        self.config_param["AO_标定电流"] = self.checkBox_36.isChecked()
        self.config_param["AO_不检测"] = self.radioButton_16.isChecked()
        self.config_param["AO_检测"] = self.radioButton_17.isChecked()
        self.config_param["AO_检测电压"] = self.checkBox_29.isChecked()
        self.config_param["AO_检测电流"] = self.checkBox_30.isChecked()
        self.config_param["AO_检测CAN"] = self.checkBox_31.isChecked()
        self.config_param["AO_检测run"] = self.checkBox_32.isChecked()
        # AI页面配置
        self.config_param["AI_型号"] = self.comboBox_3.currentIndex()
        self.config_param["AI_CAN_修改参数"] = self.checkBox_4.isChecked()
        self.config_param["AI_CAN_工装"] = self.lineEdit_18.text()
        self.config_param["AI_CAN_待检"] = self.lineEdit_19.text()
        self.config_param["AI_CAN_继电器"] = self.lineEdit_23.text()
        self.config_param["AI_附加_修改参数"] = self.checkBox_3.isChecked()
        self.config_param["AI_附加_波特率"] = self.lineEdit_16.text()
        self.config_param["AI_附加_等待时间"] = self.lineEdit_17.text()
        self.config_param["AI_附加_接收次数"] = self.lineEdit_15.text()
        self.config_param["AI_不标定"] = self.radioButton.isChecked()
        self.config_param["AI_仅标定1"] = self.radioButton_2.isChecked()
        self.config_param["AI_标定all"] = self.radioButton_3.isChecked()
        self.config_param["AI_标定电压"] = self.checkBox_5.isChecked()
        self.config_param["AI_标定电流"] = self.checkBox_6.isChecked()
        self.config_param["AI_不检测"] = self.radioButton_4.isChecked()
        self.config_param["AI_检测"] = self.radioButton_5.isChecked()
        self.config_param["AI_检测电压"] = self.checkBox_9.isChecked()
        self.config_param["AI_检测电流"] = self.checkBox_10.isChecked()
        self.config_param["AI_检测CAN"] = self.checkBox_8.isChecked()
        self.config_param["AI_检测run"] = self.checkBox_7.isChecked()
        # DIDO页面配置
        self.config_param["DIDO_型号"] = self.comboBox.currentIndex()
        self.config_param["DIDO_CAN_修改参数"] = self.checkBox_27.isChecked()
        self.config_param["DIDO_CAN_工装"] = self.lineEdit_6.text()
        self.config_param["DIDO_CAN_待检"] = self.lineEdit_7.text()
        self.config_param["DIDO_附加_修改参数"] = self.checkBox_28.isChecked()
        self.config_param["DIDO_附加_波特率"] = self.lineEdit_13.text()
        self.config_param["DIDO_附加_间隔时间"] = self.lineEdit_12.text()
        self.config_param["DIDO_附加_循环次数"] = self.lineEdit_14.text()
        self.config_param["DIDO_检测CAN"] = self.checkBox.isChecked()
        self.config_param["DIDO_检测run"] = self.checkBox_2.isChecked()

        # CPU页面配置
        for i in range(len(self.CPU_lineEdit_array)):
            self.config_param[self.CPU_lineEditName_array[i]] = self.CPU_lineEdit_array[i].text()
        self.config_param["CPU_型号"] = self.comboBox_20.currentIndex()
        self.config_param["CPU_232COM"] = self.comboBox_21.currentIndex()
        self.config_param["CPU_485COM"] = self.comboBox_22.currentIndex()
        self.config_param["CPU_typecCOM"] = self.comboBox_23.currentIndex()
        self.config_param["选项板232"] = self.comboBox_24.currentIndex()
        self.config_param["选项板485"] = self.comboBox_25.currentIndex()
        self.config_param["可编程电源"] = self.comboBox_power.currentIndex()
        for i in range(len(self.CPU_checkBox_array)):
            self.config_param[self.CPU_checkBoxName_array[i]] = self.CPU_checkBox_array[i].isChecked()

        # save
        config_str = str(self.config_param)
        config_str1 = ''
        pos = 0
        while pos >= 0:
            pos = config_str.find(',')
            config_str1 = config_str1 + config_str[:pos + 1] + '\n'
            config_str = config_str[pos + 1:]
        config_str = config_str1 + config_str
        self.configFile = open(f'{self.current_dir}/config.txt', 'w', encoding='utf-8')
        self.configFile.write(config_str)
        self.configFile.close()

    def changePN(self):
        if self.tabIndex == 0:
            if self.comboBox.currentIndex() == 1:
                self.lineEdit_PN.setText('P0000010391877')
            elif self.comboBox.currentIndex() == 4:
                self.lineEdit_PN.setText('P0000010392086')
            elif self.comboBox.currentIndex() == 5:
                self.lineEdit_PN.setText('P0000010392121')
        elif self.tabIndex == 1:
            if self.comboBox_3.currentIndex() == 0:
                self.lineEdit_PN.setText('P0000010392361')
        elif self.tabIndex == 2:
            if self.comboBox_5.currentIndex() == 0:
                self.lineEdit_PN.setText('P0000010392365')

    def changeTabWidgetPN(self):
        if self.tabIndex == 0:
            self.lineEdit.setText(self.lineEdit_PN.text())
        elif self.tabIndex == 1:
            self.lineEdit_22.setText(self.lineEdit_PN.text())
        elif self.tabIndex == 2:
            self.lineEdit_47.setText(self.lineEdit_PN.text())

    def update_time(self):
        self.label_47.setText(datetime.datetime.now().strftime('%m/%d %H:%M:%S'))
        QTimer.singleShot(1000, self.update_time)

    def killUi(self):
        self.saveConfig()
        self.close()
        try:
            os.kill(os.getpid(), signal.SIGTERM)
        except OSError as e:
            self.showInf(f"无法杀死当前进程: {e}\n")

    def label_mousePressEvent(self, event):
        if event.button() == Qt.LeftButton:
            self.dragging = True
            self.offset = event.pos()

    def label_mouseMoveEvent(self, event):
        if self.dragging and event.buttons() & Qt.LeftButton:
            self.move(event.globalPos() - self.offset)
            event.accept()

    def mouseReleaseEvent(self, event):
        self.dragging = False

    def uiResize(self):
        self.resize(740, 872)
        self.label_6.resize(740, 30)
        self.pushButton_11.setGeometry(QtCore.QRect(680, 0, 30, 30))
        self.pushButton_8.setGeometry(QtCore.QRect(710, 0, 30, 30))

        self.pushButton_9.setEnabled(False)
        self.pushButton_9.setVisible(False)
        self.pushButton_12.setEnabled(True)
        self.pushButton_12.setVisible(True)

    def uiRecovery(self):
        self.resize(1440, 872)
        self.label_6.resize(1440, 30)
        self.pushButton_11.setGeometry(QtCore.QRect(1380, 0, 30, 30))
        self.pushButton_8.setGeometry(QtCore.QRect(1410, 0, 30, 30))

        self.pushButton_12.setEnabled(False)
        self.pushButton_12.setVisible(False)
        self.pushButton_9.setEnabled(True)
        self.pushButton_9.setVisible(True)

    def setFocusLineEdit(self):
        # if len(self.lineEdit_PN.text()) == 0 and len(self.lineEdit_SN.text()) == 0 and len(self.lineEdit_REV.text()) == 0:
        #     self.lineEdit_PN.setReadOnly(False)
        #     self.lineEdit_PN.setFocus()
        if len(self.lineEdit_SN.text()) == 0 and len(self.lineEdit_REV.text()) == 0:
            self.lineEdit_SN.setReadOnly(False)
            self.lineEdit_SN.setFocus()

    def reInputPNSNREV(self):
        # self.textBrowser_5.clear()
        # self.lineEdit_PN.clear()
        self.lineEdit_SN.clear()
        self.lineEdit_REV.clear()
        # self.lineEdit_PN.setReadOnly(False)
        self.lineEdit_SN.setReadOnly(False)
        self.lineEdit_REV.setReadOnly(True)
        self.lineEdit_SN.setFocus()

        # self.lineEdit.clear()
        self.lineEdit_3.clear()
        self.lineEdit_5.clear()

        # self.lineEdit_22.clear()
        self.lineEdit_21.clear()
        self.lineEdit_20.clear()

        self.lineEdit_31.clear()
        self.lineEdit_32.clear()

        # self.lineEdit_47.clear()
        self.lineEdit_46.clear()
        self.lineEdit_45.clear()
        self.pushButton_4.setEnabled(False)
        self.pushButton_4.setStyleSheet(self.topButton_qss['off'])

    def inputPN(self):
        if len(self.lineEdit_PN.text()) == 15:
            self.lineEdit_PN.setReadOnly(True)
            self.lineEdit_SN.setReadOnly(False)
            self.lineEdit_SN.setFocus()
            if self.tabIndex == 0:
                self.lineEdit.setText(self.lineEdit_PN.text())
            elif self.tabIndex == 1:
                self.lineEdit_22.setText(self.lineEdit_PN.text())
            elif self.tabIndex == 2:
                self.lineEdit_47.setText(self.lineEdit_PN.text())

        else:
            # time.sleep(0.1)
            self.lineEdit_PN.clear()

    def inputSN(self):
        # # print(self.lineEdit_PN.text())
        if len(self.lineEdit_SN.text()) == 12:
            self.lineEdit_SN.setReadOnly(True)
            self.lineEdit_REV.setReadOnly(False)
            self.lineEdit_REV.setFocus()
            if self.tabIndex == 0:
                self.lineEdit_3.setText(self.lineEdit_SN.text())
            elif self.tabIndex == 1:
                self.lineEdit_21.setText(self.lineEdit_SN.text())
            elif self.tabIndex == 2:
                self.lineEdit_46.setText(self.lineEdit_SN.text())
            elif self.tabIndex == 3:
                self.lineEdit_31.setText(self.lineEdit_SN.text())
            elif self.tabIndex == 4:
                self.lineEdit_MA0202_SN.setText(self.lineEdit_SN.text())
        else:
            # time.sleep(0.1)
            self.lineEdit_SN.clear()

    def inputREV(self):
        if len(self.lineEdit_REV.text()) == 2:
            # reply = QMessageBox.warning(None, '操作警告', '扫入的REV码与设备中存储的REV不一致，\n取消该模块测试！',
            #                             QMessageBox.Yes | QMessageBox.No, QMessageBox.Yes)
            # if reply == QMessageBox.Yes or QMessageBox.No:
            #     self.reInputPNSNREV()
            # self.showInf('三码完成输入')
            self.lineEdit_REV.setReadOnly(True)
            self.pushButton_4.setEnabled(True)
            self.pushButton_4.setStyleSheet(self.topButton_qss['on'])
            self.pushButton_4.setStyleSheet(f"""
                                                QPushButton {{
                                                    background-color: #11826C;
                                                    color:white;
                                                    font: bold 25pt \"宋体\"
                                                }}
                                                QPushButton:hover {{
                                                    background-color: {self.hover_color.name()};
                                                    color:#212B36;
                                                    font: bold 25pt \"宋体\"
                                                }}
                                            """)
            if self.tabIndex == 0:
                self.lineEdit_5.setText(self.lineEdit_REV.text())
            elif self.tabIndex == 1:
                self.lineEdit_20.setText(self.lineEdit_REV.text())
            elif self.tabIndex == 2:
                self.lineEdit_45.setText(self.lineEdit_REV.text())
            elif self.tabIndex == 3:
                self.lineEdit_32.setText(self.lineEdit_REV.text())
            elif self.tabIndex == 4:
                self.lineEdit_MA0202_REV.setText(self.lineEdit_REV.text())
            self.pushButton_4.setFocus()
        else:
            # time.sleep(0.5)
            self.lineEdit_REV.clear()

    def changeSaveDir(self):
        self.saveDir = self.label_41.text()
        self.saveDir = QFileDialog.getExistingDirectory(self, '修改路径', self.saveDir)
        if self.saveDir != '':
            self.label_41.setText(self.saveDir)
            # print(self.saveDir)

    def openSaveDir(self):
        os.startfile(self.label_41.text())

    def tabChange(self):
        self.tabIndex = self.tabWidget.currentIndex()
        # print("current tabIndex:",self.tabIndex)
        if self.tabIndex == 0 and not self.isAllScreen:
            self.tableWidget_DIDO.setVisible(True)
            self.tableWidget_AI.setVisible(False)
            self.tableWidget_AO.setVisible(False)
            self.tableWidget_CPU.setVisible(False)
            self.label_7.setText("DIDO")
            if self.comboBox.currentIndex() == 1:
                self.lineEdit_PN.setText('P0000010391877')
            elif self.comboBox.currentIndex() == 4:
                self.lineEdit_PN.setText('P0000010392086')
            elif self.comboBox.currentIndex() == 5:
                self.lineEdit_PN.setText('P0000010392121')
        elif self.tabIndex == 1 and not self.isAllScreen:
            self.tableWidget_DIDO.setVisible(False)
            self.tableWidget_AI.setVisible(True)
            self.tableWidget_AO.setVisible(False)
            self.tableWidget_CPU.setVisible(False)
            self.label_7.setText("AI")
            if self.comboBox_3.currentIndex() == 0:
                self.lineEdit_PN.setText('P0000010392361')
        elif self.tabIndex == 2 and not self.isAllScreen:
            self.tableWidget_DIDO.setVisible(False)
            self.tableWidget_AI.setVisible(False)
            self.tableWidget_AO.setVisible(True)
            self.tableWidget_CPU.setVisible(False)
            self.label_7.setText("AO")
            if self.comboBox_5.currentIndex() == 0:
                self.lineEdit_PN.setText('P0000010392365')
        elif self.tabIndex == 3 and not self.isAllScreen:
            self.tableWidget_DIDO.setVisible(False)
            self.tableWidget_AI.setVisible(False)
            self.tableWidget_AO.setVisible(False)
            self.tableWidget_CPU.setVisible(True)
            self.label_7.setText("CPU")
            if self.comboBox_20.currentIndex() == 0:
                self.lineEdit_PN.setText('P0000010390631')
        elif self.tabIndex == 4 and not self.isAllScreen:
            self.tableWidget_DIDO.setVisible(False)
            self.tableWidget_AI.setVisible(False)
            self.tableWidget_AO.setVisible(False)
            self.tableWidget_CPU.setVisible(True)
            self.label_7.setText("CPU")
            if self.comboBox_20.currentIndex() == 0:
                self.lineEdit_PN.setText('MA0202')
        # self.reInputPNSNREV()

    def start_button(self):
        self.label.setStyleSheet(self.testState_qss['testing'])
        self.label.setText('检测中……')
        self.pushButton_3.setEnabled(True)
        self.pushButton_3.setStyleSheet(self.topButton_qss['stop_on'])
        self.pushButton_3.setStyleSheet(f"""
                                            QPushButton {{
                                                {self.topButton_qss['stop_on']}
                                            }}
                                            QPushButton:hover {{
                                                background-color: rgb(255,0,0);
                                                color:rgb(0, 0, 0)
                                            }}
                                        """)
        self.pushButton_4.setEnabled(False)
        self.pushButton_4.setVisible(False)

        self.pushButton_pause.setStyleSheet(f"""
                                                QPushButton {{
                                                    {self.topButton_qss['pause_on']}
                                                }}
                                                QPushButton:hover {{
                                                    background-color: rgb(255, 255, 0);
                                                    color:rgb(0, 0, 0)
                                                }}
                                            """)
        self.pushButton_pause.setVisible(True)
        self.pushButton_pause.setEnabled(True)
        self.textBrowser_5.clear()
        global isMainRunning
        isMainRunning = True

        if not self.startTest():
            self.PassOrFail(False)

    def pause_button(self):
        CAN_option.receivePause()
        # self.showInf(self.HORIZONTAL_LINE + '暂停中…………' + self.HORIZONTAL_LINE)
        self.label.setText('暂停中……')
        self.pushButton_pause.setVisible(False)
        self.pushButton_pause.setEnabled(False)
        self.pushButton_resume.setStyleSheet(self.topButton_qss['start_on'])
        self.pushButton_resume.setStyleSheet(f"""
                                QPushButton {{
                                    {self.topButton_qss['start_on']}
                                }}
                                QPushButton:hover {{
                                    background-color: rgb(0, 255, 0);
                                    color:rgb(0, 0, 0)
                                }}
                            """)
        self.pushButton_resume.setEnabled(True)
        self.pushButton_resume.setVisible(True)

        # self.work_thread.pause()

    def resume_button(self):
        CAN_option.receiveResume()
        self.label.setText('测试中……')
        self.pushButton_resume.setVisible(False)
        self.pushButton_resume.setEnabled(False)
        self.pushButton_pause.setStyleSheet(f"""
                                        QPushButton {{
                                            {self.topButton_qss['pause_on']}
                                        }}
                                        QPushButton:hover {{
                                            background-color: rgb(255, 255, 0);
                                            color:rgb(0, 0, 0)
                                        }}
                                    """)
        self.pushButton_pause.setEnabled(True)
        self.pushButton_pause.setVisible(True)
        # self.work_thread.bresume()

    def stop_button(self):
        # self.textBrowser_5.insertPlainText(self.HORIZONTAL_LINE + '停止测试\n' + self.HORIZONTAL_LINE)
        # self.move_to_end()
        # self.showInf(self.HORIZONTAL_LINE + '停止测试\n' + self.HORIZONTAL_LINE)
        self.stop_signal.emit(True)
        self.pushButton_3.setEnabled(False)
        self.pushButton_3.setStyleSheet(self.topButton_qss['off'])
        self.pushButton_resume.setVisible(False)
        self.pushButton_resume.setEnabled(False)
        self.pushButton_pause.setEnabled(False)
        self.pushButton_pause.setVisible(False)
        self.pushButton_4.setEnabled(False)
        self.pushButton_4.setVisible(True)
        self.pushButton_4.setStyleSheet(self.topButton_qss['off'])
        self.reInputPNSNREV()
        # self.pushButton_9.setStyleSheet(self.topButton_qss['off'])
        # self.subRunFlag = False

        CAN_option.receiveStop()

        global isMainRunning
        isMainRunning = False

    def startTest(self):
        try:
            # 启动CAN分析仪
            list_canInit = CAN_init([1])
            if not list_canInit[0]:
                self.showMessageBox_oneButton(list_canInit[1])
                self.CANFail()
                return False

            CAN_option.receiveRun()
            CAN_option.receiveResume()
            self.pushButton_10.setEnabled(False)
            self.pushButton_10.setStyleSheet('color: rgb(255, 255, 255);background-color: rgb(197, 197, 197);')
            try:
                if not self.sendMessage():  # 测试参数确定
                    return False
            except Exception as e:
                self.showInf(f"sendMessageError:{e}{self.HORIZONTAL_LINE}")
                # 捕获异常并输出详细的错误信息
                traceback.print_exc()
                return False
            QApplication.processEvents()

            # 控制可编程电源自动上电
            if self.powerControl(baudRate=9600, transmitData=self.vol_24v):
                self.showInf('成功设置电压为24V。\n')
                if self.powerControl(baudRate=9600, transmitData=self.cur_2a):
                    self.showInf('成功设置电流为2A。\n')
                    if self.powerControl(baudRate=9600, transmitData=self.power_on):
                        self.showInf('成功上电，等待设备初始化。\n')
                        if self.tabIndex == 0 or self.tabIndex == 1 or self.tabIndex == 2:
                            waitTime = 3
                        elif self.tabIndex == 3 or self.tabIndex == 4:
                            waitTime = 6
                        for i in range(waitTime):
                            self.showInf(f'剩余等待时间：{waitTime-i}秒……\n')
                            # time.sleep(1)
                    else:
                        self.showInf('上电失败，测试取消。\n')
                        return False
                else:
                    self.showInf('电流设置失败，测试取消。\n')
                    return False
            else:
                self.showInf('电压设置失败，测试取消。\n')
                return False

            # 检查三码信息
            if not mainThreadRunning():
                return False
            if len(self.module_pn) == 0 or len(self.module_sn) == 0 or len(self.module_rev) == 0:
                # self.isPause()
                if not mainThreadRunning():
                    return False
                reply = QMessageBox.warning(None, '警告', '产品三码信息不全，请重新扫入！',
                                            QMessageBox.Yes | QMessageBox.No, QMessageBox.Yes)
                return False
            self.tabIndex = self.tabWidget.currentIndex()
            # 节点分配+心跳检测
            if self.tabIndex == 0:
                if not self.configCANAddr([int(self.lineEdit_6.text()), int(self.lineEdit_7.text()), '', '', '']):
                    return False
                # DI/DO心跳检测
                if not self.isModulesOnline([self.CAN1, self.CAN2], [self.module_1, self.module_2]):
                    return False
            elif self.tabIndex == 1:
                if not self.configCANAddr([int(self.lineEdit_18.text()), int(self.lineEdit_23.text()),
                                           int(self.lineEdit_23.text()) + 1, int(self.lineEdit_19.text()), '']):
                    return False
                # AI心跳检测
                if not self.isModulesOnline([self.CAN1, self.CANAddr_relay, self.CANAddr_relay + 1, self.CAN2],
                                            [self.module_1, '继电器QR0016#1', '继电器QR0016#2', self.module_2]):
                    return False
            elif self.tabIndex == 2:
                if not self.configCANAddr([int(self.lineEdit_39.text()), int(self.lineEdit_41.text()),
                                           int(self.lineEdit_41.text()) + 1, int(self.lineEdit_40.text()), '']):
                    return False
                # AO心跳检测
                if not self.isModulesOnline([self.CAN1, self.CANAddr_relay, self.CANAddr_relay + 1, self.CAN2],
                                            [self.module_1, '继电器QR0016#1', '继电器QR0016#2', self.module_2]):
                    return False
            elif self.tabIndex == 3:
                if not self.configCANAddr([int(self.lineEdit_34.text()), int(self.lineEdit_35.text()),
                                           int(self.lineEdit_36.text())]):
                    return False
                time.sleep(0.5)
                # CPU工装心跳检测
                if not self.isModulesOnline([int(self.lineEdit_34.text()), int(self.lineEdit_35.text()),
                                             int(self.lineEdit_36.text())],
                                            ['模块ET1600', '模块QN0016', '模块QN0016']):
                    return False
            elif self.tabIndex == 4:
                if self.radioButton_MA0202.isChecked():
                    if not self.configCANAddr([int(self.lineEdit_MA0202_AE.text()),
                                               int(self.lineEdit_MA0202_AQ.text())]):
                        return False
                    # MA0202工装心跳检测
                    if not self.isModulesOnline([int(self.lineEdit_MA0202_AE.text()),
                                                 int(self.lineEdit_MA0202_AQ.text())], ['模块AE0400', '模块AQ0004']):
                        return False

            self.result_queue = Queue()
            self.testFlag = ''

            # #读设备三码
            # bool_3code,code_list = get3codeFromPLC(2)
            # if bool_3code:
            #     QMessageBox.about(None, '通过', f'三码一致！\nPN:{code_list[0]}\nSN:{code_list[1]}\nREV:{code_list[2]}')
            # else:
            #     QMessageBox.critical(None, '警告', f'三码不一致！\nPN:{code_list[0]}\nSN:{code_list[1]}\nREV:{code_list[2]}',
            #                          QMessageBox.AcceptRole | QMessageBox.RejectRole,
            #                          QMessageBox.AcceptRole)
            #     return False

            # 开始时间
            self.allStart_time = time.time()
            if not mainThreadRunning():
                return False

            self.pushButton_pause.setEnabled(True)
            self.pushButton_pause.setVisible(True)

            self.pushButton_resume.setEnabled(False)
            self.pushButton_resume.setVisible(False)

            if self.tabWidget.currentIndex() == 0:
                self.testFlag = 'DIDO'
                # self.appearanceTest(self.testFlag)
                mTable = self.tableWidget_DIDO
                if self.comboBox.currentIndex() == 0 or self.comboBox.currentIndex() == 1 \
                        or self.comboBox.currentIndex() == 2:
                    self.testFlag = 'DI'

                elif self.comboBox.currentIndex() == 3 or self.comboBox.currentIndex() == 4 \
                        or self.comboBox.currentIndex() == 5 or self.comboBox.currentIndex() == 6:
                    self.testFlag = 'DO'

                self.DIDO_thread = QThread()
                from thread_DIDO import DIDOThread
                self.DIDO_option = DIDOThread(self.inf_DIDOlist, self.result_queue, self.appearance, self.testFlag)
                self.DIDO_option.result_signal.connect(self.showInf)
                self.DIDO_option.item_signal.connect(self.DIDO_itemOperation)
                self.DIDO_option.pass_signal.connect(self.PassOrFail)
                # self.DIDO_option.RunErr_signal.connect(self.testRunErr)
                # self.DIDO_option.CANRunErr_signal.connect(self.testCANRunErr)
                self.DIDO_option.messageBox_signal.connect(self.showMessageBox)
                self.DIDO_option.pic_messageBox_signal.connect(self.pic_MessageBox)
                self.DIDO_option.messageBox_oneButton_signal.connect(self.showMessageBox_oneButton)
                self.DIDO_option.gif_messageBox_signal.connect(self.gif_MessageBox)
                self.DIDO_option.checkBox_messageBox_signal.connect(self.generateCheckBox_messageBox_widget)#widget在QMessageBox关闭后也会消失，因此需要每次重新生成新的。
                self.DIDO_option.checkBox_messageBox_signal.connect(self.checkBox_MessageBox)
                # self.DIDO_option.excel_signal.connect(self.generateExcel)
                self.DIDO_option.allFinished_signal.connect(self.allFinished)
                self.DIDO_option.label_signal.connect(self.labelChange)
                self.DIDO_option.saveExcel_signal.connect(self.saveExcel)
                # self.DIDO_option.print_signal.connect(self.printResult)

                self.pushButton_3.clicked.connect(self.DIDO_option.stop_work)
                self.pushButton_pause.clicked.connect(self.DIDO_option.pause_work)
                self.pushButton_resume.clicked.connect(self.DIDO_option.resume_work)

                self.DIDO_option.moveToThread(self.DIDO_thread)
                self.DIDO_thread.started.connect(self.DIDO_option.DIDOOption)
                self.DIDO_thread.start()

            if self.tabWidget.currentIndex() == 1:
                # self.lineEdit_PN = 'PRDr9HPA06Mz-00'
                # self.lineEdit_SN = 'S1223-001083'
                # self.lineEdit_REV = '06'

                self.testFlag = 'AI'
                # self.appearanceTest(self.testFlag)
                # self.appearance = True
                # self.showInf(f'self.tabWidget.currentIndex()={self.tabWidget.currentIndex()}\n\n')
                mTable = self.tableWidget_AI
                # isPassVol = True
                # isPassCur = True
                if self.isTest:
                    if self.isTestRunErr:
                        self.testNum = self.testNum - 1
                        # self.testRunErr(self.CANAddr_AI)
                    elif not self.isTestRunErr:
                        if not mainThreadRunning():
                            return False
                        self.AI_itemOperation([1, 0, 0, ''])
                        if not mainThreadRunning():
                            return False
                        self.AI_itemOperation([2, 0, 0, ''])
                    # self.showInf(f'RunErrself.testNum = {self.testNum}\n\n')
                    if self.isTestCANRunErr:
                        self.testNum = self.testNum - 1
                        # self.testCANRunErr(self.CANAddr_AI)
                    elif not self.isTestCANRunErr:
                        if not mainThreadRunning():
                            return False
                        self.AI_itemOperation([3, 0, 0, ''])
                        if not mainThreadRunning():
                            return False
                        self.AI_itemOperation([4, 0, 0, ''])
                        # self.itemOperation(mTable, 3, 0, 0, '')
                        # self.itemOperation(mTable, 4, 0, 0, '')
                    if self.isAITestVol:
                        self.testNum = self.testNum - 1
                    elif not self.isAITestVol:
                        if not mainThreadRunning():
                            return False
                        self.AI_itemOperation([7, 0, 0, ''])
                    if self.isAITestCur:
                        self.testNum = self.testNum - 1
                    elif not self.isAITestCur:
                        if not mainThreadRunning():
                            return False
                        self.AI_itemOperation([8, 0, 0, ''])

                # if self.isCalibrate:
                # self.calibrateAI(mTable)
                self.AI_thread = None
                self.worker = None
                if not self.AI_thread or not self.worker_thread.isRunning():
                    # # 创建队列用于主线程和子线程之间的通信
                    # self.result_queue = Queue()

                    self.AI_thread = QThread()
                    from thread_AI import AIThread
                    self.AI_option = AIThread(self.inf_AIlist, self.result_queue, self.appearance)

                    self.AI_option.result_signal.connect(self.showInf)
                    self.AI_option.item_signal.connect(self.AI_itemOperation)
                    self.AI_option.pass_signal.connect(self.PassOrFail)
                    # self.AI_option.RunErr_signal.connect(self.testRunErr)
                    # self.AI_option.CANRunErr_signal.connect(self.testCANRunErr)
                    self.AI_option.messageBox_signal.connect(self.showMessageBox)
                    self.AI_option.pic_messageBox_signal.connect(self.pic_MessageBox)
                    self.AI_option.messageBox_oneButton_signal.connect(self.showMessageBox_oneButton)
                    # self.AI_option.excel_signal.connect(self.generateExcel)
                    self.AI_option.allFinished_signal.connect(self.allFinished)
                    self.AI_option.label_signal.connect(self.labelChange)
                    self.AI_option.saveExcel_signal.connect(self.saveExcel)  # 保存测试报告
                    # self.AI_option.print_signal.connect(self.printResult)#打印测试标签

                    self.pushButton_3.clicked.connect(self.AI_option.stop_work)
                    self.pushButton_pause.clicked.connect(self.AI_option.pause_work)
                    self.pushButton_resume.clicked.connect(self.AI_option.resume_work)

                    self.AI_option.moveToThread(self.AI_thread)
                    self.AI_thread.started.connect(self.AI_option.AIOption)
                    self.AI_thread.start()



            elif self.tabWidget.currentIndex() == 2:
                self.testFlag = 'AO'
                mTable = self.tableWidget_AO

                # self.appearanceTest(self.testFlag)
                self.AO_thread = None
                self.worker = None
                if not self.AO_thread or not self.worker_thread.isRunning():
                    # # 创建队列用于主线程和子线程之间的通信
                    # self.result_queue = Queue()

                    self.AO_thread = QThread()
                    from thread_AO import AOThread
                    self.AO_option = AOThread(self.inf_AOlist, self.result_queue, self.appearance)
                    self.AO_option.result_signal.connect(self.showInf)
                    self.AO_option.item_signal.connect(self.AO_itemOperation)
                    self.AO_option.pass_signal.connect(self.PassOrFail)
                    # self.AO_option.RunErr_signal.connect(self.testRunErr)
                    # self.AO_option.CANRunErr_signal.connect(self.testCANRunErr)
                    self.AO_option.messageBox_signal.connect(self.showMessageBox)
                    self.AO_option.pic_messageBox_signal.connect(self.pic_MessageBox)
                    self.AO_option.messageBox_oneButton_signal.connect(self.showMessageBox_oneButton)
                    # self.AO_option.excel_signal.connect(self.generateExcel)
                    self.AO_option.allFinished_signal.connect(self.allFinished)
                    self.AO_option.label_signal.connect(self.labelChange)
                    self.AO_option.saveExcel_signal.connect(self.saveExcel)
                    # self.AO_option.print_signal.connect(self.printResult)

                    self.pushButton_3.clicked.connect(self.AO_option.stop_work)
                    self.pushButton_pause.clicked.connect(self.AO_option.pause_work)
                    self.pushButton_resume.clicked.connect(self.AO_option.resume_work)

                    self.AO_option.moveToThread(self.AO_thread)
                    self.AO_thread.started.connect(self.AO_option.AOOption)
                    self.AO_thread.start()

            elif self.tabWidget.currentIndex() == 3:
                self.testFlag = 'CPU'
                mTable = self.tableWidget_CPU
                # CPU的外观检测选项放在子线程里进行并且可选
                # self.appearanceTest(self.testFlag)
                self.CPU_thread = None
                self.worker = None
                if not self.CPU_thread or not self.worker_thread.isRunning():
                    # # 创建队列用于主线程和子线程之间的通信
                    # self.result_queue = Queue()

                    self.CPU_thread = QThread()
                    from thread_CPU import CPUThread
                    self.CPU_option = CPUThread(self.inf_CPUlist, self.result_queue)
                    self.pushButton_3.clicked.connect(self.CPU_option.cancelAllTest)
                    self.CPU_option.result_signal.connect(self.showInf)
                    self.CPU_option.item_signal.connect(self.CPU_itemOperation)
                    self.CPU_option.pass_signal.connect(self.PassOrFail)
                    # self.CPU_option.RunErr_signal.connect(self.testRunErr)
                    # self.CPU_option.CANRunErr_signal.connect(self.testCANRunErr)
                    self.CPU_option.messageBox_signal.connect(self.showMessageBox)
                    self.CPU_option.pic_messageBox_signal.connect(self.pic_MessageBox)
                    self.CPU_option.messageBox_oneButton_signal.connect(self.showMessageBox_oneButton)
                    self.CPU_option.moveToRow_signal.connect(self.CPU_moveToRow)
                    # self.CPU_option.excel_signal.connect(self.generateExcel)
                    self.CPU_option.allFinished_signal.connect(self.allFinished)
                    self.CPU_option.label_signal.connect(self.labelChange)
                    self.CPU_option.saveExcel_signal.connect(self.saveExcel)
                    # self.CPU_option.print_signal.connect(self.printResult)

                    self.pushButton_3.clicked.connect(self.CPU_option.stop_work)
                    self.pushButton_pause.clicked.connect(self.CPU_option.pause_work)
                    self.pushButton_resume.clicked.connect(self.CPU_option.resume_work)

                    self.CPU_option.moveToThread(self.CPU_thread)
                    self.CPU_thread.started.connect(self.CPU_option.CPU_start)
                    self.CPU_thread.start()

            elif self.tabWidget.currentIndex() == 4:
                self.testFlag = 'CPU'
                mTable = self.tableWidget_CPU
                # CPU的外观检测选项放在子线程里进行并且可选
                # self.appearanceTest(self.testFlag)
                self.CPU_thread = None
                self.worker = None
                if not self.CPU_thread or not self.worker_thread.isRunning():
                    # # 创建队列用于主线程和子线程之间的通信
                    # self.result_queue = Queue()
                    self.inf_MA0202_test = [False for x in range(18)]
                    self.inf_MA0202_test[17] = True
                    self.inf_CPUlist = [['', '', '', '', self.moduleName_1, self.moduleName_2, 1],
                                        [self.radioButton_MA0202.text(), self.lineEdit_MA0202_PN.text(),
                                         self.lineEdit_MA0202_SN.text(), self.lineEdit_MA0202_REV.text(), '', 0, 0],
                                        ['', '', '', int(self.lineEdit_MA0202_AE.text()),
                                         int(self.lineEdit_MA0202_AQ.text())],
                                        ['', '', self.comboBox_MA0202_typeC.currentText(), self.saveDir, '', '', ''],
                                        [self.inf_MA0202_test, ''],
                                        self.current_dir, 'MA0202']
                    self.CPU_thread = QThread()
                    from thread_CPU import CPUThread
                    self.CPU_option = CPUThread(self.inf_CPUlist, self.result_queue)
                    self.pushButton_3.clicked.connect(self.CPU_option.cancelAllTest)
                    self.CPU_option.result_signal.connect(self.showInf)
                    self.CPU_option.item_signal.connect(self.CPU_itemOperation)
                    self.CPU_option.pass_signal.connect(self.PassOrFail)
                    # self.CPU_option.RunErr_signal.connect(self.testRunErr)
                    # self.CPU_option.CANRunErr_signal.connect(self.testCANRunErr)
                    self.CPU_option.messageBox_signal.connect(self.showMessageBox)
                    self.CPU_option.pic_messageBox_signal.connect(self.pic_MessageBox)
                    self.CPU_option.messageBox_oneButton_signal.connect(self.showMessageBox_oneButton)
                    self.CPU_option.moveToRow_signal.connect(self.CPU_moveToRow)
                    # self.CPU_option.excel_signal.connect(self.generateExcel)
                    self.CPU_option.allFinished_signal.connect(self.allFinished)
                    self.CPU_option.label_signal.connect(self.labelChange)
                    self.CPU_option.saveExcel_signal.connect(self.saveExcel)
                    # self.CPU_option.print_signal.connect(self.printResult)

                    self.pushButton_3.clicked.connect(self.CPU_option.stop_work)
                    self.pushButton_pause.clicked.connect(self.CPU_option.pause_work)
                    self.pushButton_resume.clicked.connect(self.CPU_option.resume_work)

                    self.CPU_option.moveToThread(self.CPU_thread)
                    self.CPU_thread.started.connect(self.CPU_option.CPU_start)
                    self.CPU_thread.start()
        except Exception as e:
            self.showInf(f"startTestError:{e}{self.HORIZONTAL_LINE}")
            # 捕获异常并输出详细的错误信息
            traceback.print_exc()
            return False

        return True

    def powerControl(self,baudRate:int,transmitData:list):
        isReceiveTrueData:bool = True
        try:
            m_port = ' '
            #打开串口
            m_port = self.comboBox_power.currentText()
            m_serial = serial.Serial(port=m_port, baudrate=baudRate, timeout=1)
            for i in range(5):#每条指令发送5遍，确保能正常控制可编程电源
                # 发送数据
                m_serial.write(bytes(transmitData))
                time.sleep(0.1)
        except serial.SerialException as e:
            self.showMessageBox_oneButton(['错误警告', f'串口{m_port}打开失败，请检查该串口是否被占用。'])
            isReceiveTrueData &= False
            # 关闭串口
            m_serial.close()
            # return isReceiveTrueData
        finally:
            # 关闭串口
            m_serial.close()
            return isReceiveTrueData
    def appearanceTest(self, testFlag):
        if testFlag == 'DIDO':
            mTable = self.tableWidget_DIDO
        elif testFlag == 'AI':
            mTable = self.tableWidget_AI
        elif testFlag == 'AO':
            mTable = self.tableWidget_AO
        # elif testFlag == 'CPU':
        #     mTable = self.tableWidget_CPU
        # 检测外观
        appearanceStart_time = time.time()
        if not mainThreadRunning():
            return False

        self.itemOperation(mTable, 0, 1, 0, '')

        if not mainThreadRunning():
            return False
        reply = QMessageBox.question(None, '外观检测',
                                     '请检查：\n（1）外壳字体是否清晰?\n（2）型号是否正确？\n（3）外壳是否完好？',
                                     QMessageBox.Yes | QMessageBox.No, QMessageBox.Yes)
        if reply == QMessageBox.Yes:
            self.appearance = True
            appearanceEnd_time = time.time()
            appearanceTest_time = round(appearanceEnd_time - appearanceStart_time, 2)
            if not mainThreadRunning():
                return False
            self.itemOperation(mTable, 0, 2, 1, appearanceTest_time)

        elif reply == QMessageBox.No:
            self.appearance = False
            appearanceEnd_time = time.time()
            appearanceTest_time = round(appearanceEnd_time - appearanceStart_time, 2)
            if not mainThreadRunning():
                return False
            self.itemOperation(mTable, 0, 2, 2, appearanceTest_time)

        self.testNum = self.testNum - 1

        return True

    def labelChange(self, list):
        self.label.setStyleSheet(self.testState_qss[list[0]])
        self.label.setText(list[1])

    def allFinished(self):
        end_time = time.time()
        total_time = round(end_time - self.allStart_time, 3)
        self.showInf(f'本轮测试总时间：{total_time} 秒' + self.HORIZONTAL_LINE)
        # 关闭CAN分析仪
        # CAN_option.close(CAN_option.VCI_USB_CAN_2, CAN_option.DEV_INDEX)
        # self.lineEdit_PN.clear()
        self.lineEdit_SN.clear()
        self.lineEdit_REV.clear()
        # self.lineEdit_PN.setReadOnly(False)
        # self.lineEdit_SN.setReadOnly(False)
        # self.lineEdit_REV.setReadOnly(False)
        if self.tabWidget.currentIndex() == 0:
            # self.lineEdit.clear()
            self.lineEdit_3.clear()
            self.lineEdit_5.clear()
        elif self.tabWidget.currentIndex() == 1:
            # self.lineEdit_22.clear()
            self.lineEdit_21.clear()
            self.lineEdit_20.clear()
        elif self.tabWidget.currentIndex() == 2:
            self.lineEdit_45.clear()
            self.lineEdit_46.clear()
            # self.lineEdit_47.clear()
        elif self.tabWidget.currentIndex() == 3:
            self.lineEdit_31.clear()
            self.lineEdit_32.clear()
            # self.lineEdit_47.clear()
        self.lineEdit_SN.setFocus()
        self.endOfTest()

    def showMessageBox(self, list):
        # # 在主线程中创建消息框并显示
        # reply = QMessageBox.question(None, list[0], list[1],
        #                              QMessageBox.AcceptRole | QMessageBox.RejectRole,
        #                              QMessageBox.AcceptRole)
        #
        # # 将弹窗结果放入队列
        # self.result_queue.put(reply)
        global isMainRunning
        if isMainRunning:
            msg_box = QMessageBox()
            # 设置消息
            msg_box.setText(list[1])
            # msg_box.setInformativeText(list[1])
            msg_box.setWindowTitle(list[0])
            # 设置样式表
            msg_box.setStyleSheet('QLabel{font-size: 18px;}')
            # 添加按钮
            msg_box.addButton('是', QMessageBox.AcceptRole)
            msg_box.addButton('否', QMessageBox.RejectRole)

            # 显示消息框
            reply = msg_box.exec_()
            # 将弹窗结果放入队列
            self.result_queue.put(reply)
        else:
            # 将弹窗结果放入队列
            self.result_queue.put(QMessageBox.RejectRole)
    def showMessageBox_oneButton(self, list):
        # # 在主线程中创建消息框并显示
        # reply = QMessageBox.question(None, list[0], list[1],
        #                              QMessageBox.AcceptRole | QMessageBox.RejectRole,
        #                              QMessageBox.AcceptRole)
        #
        # # 将弹窗结果放入队列
        # self.result_queue.put(reply)
        global isMainRunning
        if isMainRunning:
            msg_box = QMessageBox()
            # 设置消息
            msg_box.setText(list[1])
            # msg_box.setInformativeText(list[1])
            msg_box.setWindowTitle(list[0])
            # 设置样式表
            msg_box.setStyleSheet('QLabel{font-size: 18px;}')
            # 添加按钮
            msg_box.addButton('确定', QMessageBox.AcceptRole)
            # msg_box.addButton('否', QMessageBox.RejectRole)

            # 显示消息框
            reply = msg_box.exec_()
            # 将弹窗结果放入队列
            self.result_queue.put(reply)
        else:
            # 将弹窗结果放入队列
            self.result_queue.put(QMessageBox.RejectRole)

    def pic_MessageBox(self, list):
        global isMainRunning
        if isMainRunning:
            msg_box = QMessageBox()

            # 设置图标
            icon = QIcon(list[2])
            msg_box.setIconPixmap(icon.pixmap(300, 300))

            # 设置消息
            msg_box.setText(list[1])
            # msg_box.setInformativeText(list[1])
            msg_box.setWindowTitle(list[0])
            # 设置样式表
            msg_box.setStyleSheet('QLabel{font-size: 18px;}')

            # 添加按钮
            msg_box.addButton('是', QMessageBox.AcceptRole)
            msg_box.addButton('否', QMessageBox.RejectRole)

            # 显示消息框
            reply = msg_box.exec_()

            # 将弹窗结果放入队列
            self.result_queue.put(reply)
            # # 创建一个定时器，3秒后关闭提示框
            # timer = QTimer()
            # timer.setInterval(3000)  # 3000毫秒即3秒
            # timer.timeout.connect(msg_box.close)
            # timer.start()
        else:
            # 将弹窗结果放入队列
            self.result_queue.put(QMessageBox.RejectRole)

    def gif_MessageBox(self, list):
        global isMainRunning
        if isMainRunning:
            msg_box = gif_QMessageBox.MyMessageBox()
            # 设置gif图标
            msg_box.movie = QMovie(list[2])
            # 设置gif图标尺寸
            size = QSize(400, 300)
            msg_box.movie.setScaledSize(size)
            msg_box.label.setMovie(msg_box.movie)

            # 设置消息
            msg_box.label_text.setText(list[1])
            # msg_box.setText(list[1])
            msg_box.setWindowTitle(list[0])
            # 设置样式表
            msg_box.setStyleSheet('QLabel{font-size: 18px;}')

            # 添加按钮
            msg_box.addButton('是', QMessageBox.AcceptRole)
            msg_box.addButton('否', QMessageBox.RejectRole)

            # 显示消息框
            reply = msg_box.exec_()

            # # 将弹窗结果放入队列
            self.result_queue.put(reply)
        else:
            # 将弹窗结果放入队列
            self.result_queue.put(QMessageBox.RejectRole)

    def checkBox_MessageBox(self, list):
        global isMainRunning

        if isMainRunning:
            msg_box = QMessageBox()

            # 将QWidget对象添加到QMessageBox中
            msg_box.layout().addWidget(self.checkBox_messageBox_widget)

            # 设置QMessageBox的文本和标准按钮
            msg_box.setText(list[1])
            # 添加按钮
            msg_box.addButton('确定', QMessageBox.AcceptRole)
            msg_box.setWindowTitle(list[0])
            # 设置样式表
            msg_box.setStyleSheet('QLabel{font-size: 18px;}')

            # 显示消息框
            reply = msg_box.exec_()


            # # 将弹窗结果放入队列
            self.result_queue.put([reply,self.checkBox_status])
        else:
            # 将弹窗结果放入队列
            self.result_queue.put([QMessageBox.RejectRole,self.checkBox_status])

    def CPU_moveToRow(self, list):
        self.tableWidget_CPU.setCurrentCell(list[0], list[1])

    def PassOrFail(self, isPass):
        #控制可编程电源断电
        if not self.powerControl(baudRate=9600, transmitData=self.power_off):
            self.showMessageBox_oneButton(['操作提示','自动断电失败，请手动断电后再取下待测模块。'])
        if not isPass:
            self.endOfTest()
            self.initPara()
            # self.textBrowser_5.clear()
            self.errorInf = ''
            self.errorNum = 0
            self.isAIPassTest = True
            self.isAIVolPass = True
            self.isAICurPass = True

            self.isAOPassTest = True
            self.isAOVolPass = True
            self.isAOCurPass = True

            self.isLEDRunOK = True
            self.isLEDErrOK = True
            self.CAN_runLED = True
            self.CAN_errorLED = True
            self.stop_subThread()

            # self.label.setStyleSheet(self.testState_qss['stop'])
            # self.label.setText('')
        else:
            self.stop_subThread()

        self.pushButton_10.setEnabled(True)
        self.pushButton_10.setStyleSheet(f"""
                                        QPushButton {{
                                            background-color: #11826C;
                                            color:white
                                        }}
                                        QPushButton:hover {{
                                            background-color: {self.hover_color.name()};
                                            color:#212B36
                                        }}
                                    """)

    def stop_subThread(self):
        try:
            if self.testFlag == 'AI':
                if self.AI_thread and self.AI_thread.isRunning():
                    try:
                        self.AI_thread.quit()
                        self.AI_thread.wait()
                        # self.showInf(f'结束AI子线程成功！' + self.HORIZONTAL_LINE)
                    except Exception as e:
                        self.showInf(f'结束AI线程异常！' + self.HORIZONTAL_LINE)
                        traceback.print_exc()
            elif self.testFlag == 'AO':
                if self.AO_thread and self.AO_thread.isRunning():
                    try:
                        self.AO_thread.quit()
                        self.AO_thread.wait()
                    except Exception as e:
                        self.showInf(f'结束AO线程异常！' + self.HORIZONTAL_LINE)
                        traceback.print_exc()
            elif self.testFlag == 'DO' or self.testFlag == 'DI' or self.testFlag == 'DIDO':
                if self.DIDO_thread and self.DIDO_thread.isRunning():
                    # self.showInf(f'结束DIDO子线程' + self.HORIZONTAL_LINE)
                    try:
                        self.DIDO_thread.quit()
                        self.DIDO_thread.wait()
                    except Exception as e:
                        self.showInf(f'结束DIDO线程异常！' + self.HORIZONTAL_LINE)
                        traceback.print_exc()
            elif self.testFlag == 'CPU':
                if self.CPU_thread and self.CPU_thread.isRunning():
                    try:
                        # self.showInf(f'结束CPU子线程' + self.HORIZONTAL_LINE)
                        self.CPU_thread.quit()
                        self.CPU_thread.wait()
                    except Exception as e:
                        self.showInf(f'结束CPU线程异常！' + self.HORIZONTAL_LINE)
                        traceback.print_exc()
            else:
                self.showInf(f'不存在运行的子线程' + self.HORIZONTAL_LINE)

            self.textBrowser_5.append('进程已正确停止。\n')
            QApplication.processEvents()
            time.sleep(0.1)
        except:
            self.textBrowser_5.append('进程停止出错。\n')
            QApplication.processEvents()
            time.sleep(0.1)
        finally:
            self.lineEdit_SN.setReadOnly(False)
            self.lineEdit_REV.setReadOnly(False)

    def DIDOCANAddr_stateChanged(self):
        self.lineEdit_6.setEnabled(self.checkBox_27.isChecked())
        self.label_16.setEnabled(self.checkBox_27.isChecked())
        self.lineEdit_7.setEnabled(self.checkBox_27.isChecked())
        self.label_17.setEnabled(self.checkBox_27.isChecked())
        # self.saveConfig()

    def DIDOAdPara_stateChanged(self):
        self.lineEdit_13.setEnabled(self.checkBox_28.isChecked())
        self.label_22.setEnabled(self.checkBox_28.isChecked())
        self.lineEdit_12.setEnabled(self.checkBox_28.isChecked())
        self.label_23.setEnabled(self.checkBox_28.isChecked())
        self.lineEdit_14.setEnabled(self.checkBox_28.isChecked())
        self.label_26.setEnabled(self.checkBox_28.isChecked())
        self.label_25.setEnabled(self.checkBox_28.isChecked())
        self.label_24.setEnabled(self.checkBox_28.isChecked())
        # self.saveConfig()

    """

    AI

    """

    def AICANAddr_stateChanged(self):
        self.lineEdit_18.setEnabled(self.checkBox_4.isChecked())
        self.label_48.setEnabled(self.checkBox_4.isChecked())
        self.lineEdit_19.setEnabled(self.checkBox_4.isChecked())
        self.label_49.setEnabled(self.checkBox_4.isChecked())
        self.lineEdit_23.setEnabled(self.checkBox_4.isChecked())
        self.label_54.setEnabled(self.checkBox_4.isChecked())
        # self.saveConfig()

    def AIAdPara_stateChanged(self):
        self.lineEdit_15.setEnabled(self.checkBox_3.isChecked())
        self.label_42.setEnabled(self.checkBox_3.isChecked())
        self.lineEdit_16.setEnabled(self.checkBox_3.isChecked())
        self.label_43.setEnabled(self.checkBox_3.isChecked())
        self.lineEdit_17.setEnabled(self.checkBox_3.isChecked())
        self.label_44.setEnabled(self.checkBox_3.isChecked())
        self.label_45.setEnabled(self.checkBox_3.isChecked())
        self.label_46.setEnabled(self.checkBox_3.isChecked())
        # self.saveConfig()

    def CPU_paramChanged(self):
        CPU_param_array = [self.lineEdit_34, self.lineEdit_35, self.lineEdit_36,
                           self.lineEdit_37, self.lineEdit_38, self.comboBox_20, self.comboBox_21,
                           self.comboBox_22, self.comboBox_23, self.comboBox_24, self.comboBox_25,
                           self.comboBox_power, self.label_59, self.label_60, self.label_62,
                           self.label_64, self.label_65, self.label_66, self.label_67, self.label_68, self.label_69,
                           self.label_70, self.label_71, self.label_72, self.label_90, self.pushButton_CPU_renewSerial]
        for Cparam in CPU_param_array:
            Cparam.setEnabled(self.checkBox_71.isChecked())

    def MA0202_paramChanged(self):
        MA0202_param_array = [self.label_21, self.radioButton_MA0202, self.label_105,
                              self.comboBox_MA0202_typeC, self.pushButton_MA0202_renewSerial, self.label_116,
                              self.label_113,
                              self.lineEdit_MA0202_AE, self.label_114, self.lineEdit_MA0202_AQ]
        for Mparam in MA0202_param_array:
            Mparam.setEnabled(self.checkBox_MA0202_para.isChecked())

    def testAllorNot(self):
        CPU_test_array = [self.checkBox_50, self.checkBox_51, self.checkBox_52, self.checkBox_53,
                          self.checkBox_55, self.checkBox_56, self.checkBox_57,
                          self.checkBox_58, self.checkBox_59, self.checkBox_60, self.checkBox_61,
                          self.checkBox_62, self.checkBox_63, self.checkBox_64, self.checkBox_65,
                          self.checkBox_66, self.checkBox_54, self.checkBox_67
                          ]
        if self.checkBox_73.isChecked():
            self.checkBox_73.setText("取消全选")
            for i in range(len(CPU_test_array)):
                CPU_test_array[i].setChecked(True)
        else:
            self.checkBox_73.setText("全选")
            for i in range(len(CPU_test_array)):
                CPU_test_array[i].setChecked(False)

    def AI_NotCalibrate(self):
        self.checkBox_5.setEnabled(False)
        self.checkBox_6.setEnabled(False)

    def AI_Calibrate(self):
        self.checkBox_5.setEnabled(True)
        self.checkBox_6.setEnabled(True)

    def AI_NotTest(self):
        self.checkBox_7.setEnabled(False)
        self.checkBox_8.setEnabled(False)
        self.checkBox_9.setEnabled(False)
        self.checkBox_10.setEnabled(False)

    def AI_Test(self):
        self.checkBox_7.setEnabled(True)
        self.checkBox_8.setEnabled(True)
        self.checkBox_9.setEnabled(True)
        self.checkBox_10.setEnabled(True)

    """

    AO

    """

    def AOCANAddr_stateChanged(self):
        self.lineEdit_39.setEnabled(self.checkBox_33.isChecked())
        self.label_74.setEnabled(self.checkBox_33.isChecked())
        self.lineEdit_40.setEnabled(self.checkBox_33.isChecked())
        self.label_75.setEnabled(self.checkBox_33.isChecked())
        self.lineEdit_41.setEnabled(self.checkBox_33.isChecked())
        self.label_76.setEnabled(self.checkBox_33.isChecked())
        # self.saveConfig()

    def AOAdPara_stateChanged(self):
        self.lineEdit_42.setEnabled(self.checkBox_34.isChecked())
        self.label_77.setEnabled(self.checkBox_34.isChecked())
        self.lineEdit_43.setEnabled(self.checkBox_34.isChecked())
        self.label_78.setEnabled(self.checkBox_34.isChecked())
        self.lineEdit_44.setEnabled(self.checkBox_34.isChecked())
        self.label_79.setEnabled(self.checkBox_34.isChecked())
        self.label_80.setEnabled(self.checkBox_34.isChecked())
        self.label_81.setEnabled(self.checkBox_34.isChecked())
        # self.saveConfig()

    def AO_NotCalibrate(self):
        self.checkBox_35.setEnabled(False)
        self.checkBox_36.setEnabled(False)

    def AO_Calibrate(self):
        self.checkBox_35.setEnabled(True)
        self.checkBox_36.setEnabled(True)

    def AO_NotTest(self):
        self.checkBox_29.setEnabled(False)
        self.checkBox_30.setEnabled(False)
        self.checkBox_31.setEnabled(False)
        self.checkBox_32.setEnabled(False)

    def AO_Test(self):
        self.checkBox_29.setEnabled(True)
        self.checkBox_30.setEnabled(True)
        self.checkBox_31.setEnabled(True)
        self.checkBox_32.setEnabled(True)

    def allScreen(self):
        self.label_7.setVisible(False)
        self.isAllScreen = True
        for tW in self.tableWidget_array:
            tW.setVisible(False)
        # self.tableWidget_AI.setVisible(False)
        # self.tableWidget_AO.setVisible(False)
        self.pushbutton_allScreen.setEnabled(False)
        self.pushbutton_allScreen.setVisible(False)
        self.label_28.setGeometry(10, 210, 192, 30)
        self.textBrowser_5.setMaximumSize(711, 800)
        self.textBrowser_5.setGeometry(10, 240, 711, 601)

        self.pushbutton_cancelAllScreen.setEnabled(True)
        self.pushbutton_cancelAllScreen.setVisible(True)
        QApplication.processEvents()

    def cancelAllScreen(self):
        self.label_7.setVisible(True)
        self.isAllScreen = False
        self.pushbutton_cancelAllScreen.setEnabled(False)
        self.pushbutton_cancelAllScreen.setVisible(False)

        self.textBrowser_5.setGeometry(10, 664, 711, 177)
        self.label_28.setGeometry(10, 634, 192, 30)
        self.tabIndex = self.tabWidget.currentIndex()

        self.pushbutton_allScreen.setEnabled(True)
        self.pushbutton_allScreen.setVisible(True)
        if self.tabIndex == 0:
            self.tableWidget_DIDO.setVisible(True)
        elif self.tabIndex == 1:
            self.tableWidget_AI.setVisible(True)
        elif self.tabIndex == 2:
            self.tableWidget_AO.setVisible(True)
        elif self.tabIndex == 3:
            self.tableWidget_CPU.setVisible(True)
        QApplication.processEvents()

    def sendMessage(self):
        self.textBrowser_5.clear()
        # 清空ASCII码列表
        self.asciiCode_pn = []
        self.asciiCode_sn = []
        self.asciiCode_rev = []

        self.errorInf = ''
        self.errorNum = 0
        self.isAIPassTest = True
        self.isAIVolPass = True
        self.isAICurPass = True

        self.isAOPassTest = True
        self.isAOVolPass = True
        self.isAOCurPass = True

        # DI测试是否通过
        isDIPassTest = True
        # DO测试是否通过
        isDOPassTest = True

        self.isLEDRunOK = True
        self.isLEDErrOK = True
        self.CAN_runLED = True
        self.CAN_errorLED = True

        self.tabIndex = self.tabWidget.currentIndex()
        if self.tabIndex == 0:  # DI/DO界面
            mTable = self.tableWidget_DIDO
            # print(f'tabIndex={self.tabIndex}')

            self.testNum = 2  # 外观检测 + CAN_RunErr检测 + RunErr检测 + 通道检测
            # 获取产品信息
            self.module_type = self.comboBox.currentText()
            if self.module_type == 'ET0800' or self.module_type == 'ET1600' or \
                    self.module_type == 'ET3200':
                self.m_Channels = int(self.module_type[2:4])
            elif self.module_type == 'QN0008' or self.module_type == 'QN0016' or \
                    self.module_type == 'QN0032':
                self.m_Channels = int(self.module_type[4:])

            self.showInf(f'模块类型：{self.module_type}，'
                         f'通道：{self.m_Channels}' + self.HORIZONTAL_LINE)
            self.module_pn = self.lineEdit.text()
            self.module_sn = self.lineEdit_3.text()
            self.module_rev = self.lineEdit_5.text()
            # print(f'{self.module_type},{self.module_pn},{self.module_sn},{self.module_rev}')

            if self.comboBox.currentIndex() == 0 or self.comboBox.currentIndex() == 1 \
                    or self.comboBox.currentIndex() == 2:
                self.module_1 = '模块QP(N)0016'
                self.module_2 = '模块ET1600'
                self.m_Channels = int(self.module_type[2:4])
                # self.inf_param = [mTable, self.module_1, self.module_2, self.testNum]
                # 获取CAN地址
                self.CANAddr_DI = int(self.lineEdit_7.text())
                self.CANAddr_DO = int(self.lineEdit_6.text())
                self.CAN1 = self.CANAddr_DO
                self.CAN2 = self.CANAddr_DI

            if self.comboBox.currentIndex() == 3 or self.comboBox.currentIndex() == 4 \
                    or self.comboBox.currentIndex() == 5:
                self.module_1 = '模块ET1600'
                if self.comboBox.currentIndex() == 3:
                    self.module_2 = '模块QN0008'
                elif self.comboBox.currentIndex() == 4:
                    self.module_2 = '模块QN0016'
                elif self.comboBox.currentIndex() == 5:
                    self.module_2 = '模块QP0016'
                self.m_Channels = int(self.module_type[4:])
                # self.inf_param = [mTable, self.module_1, self.module_2, self.testNum]

                # 获取CAN地址
                self.CANAddr_DI = int(self.lineEdit_6.text())
                self.CANAddr_DO = int(self.lineEdit_7.text())
                self.CAN1 = self.CANAddr_DI
                self.CAN2 = self.CANAddr_DO
            # 获取附加参数
            self.baud_rate = int(self.lineEdit_13.text())
            self.waiting_time = int(self.lineEdit_12.text())
            self.loop_num = int(self.lineEdit_14.text())
            self.saveDir = self.label_41.text()  # 保存路径
            # # print(f'{self.module_type},{self.module_pn},{self.module_sn},{self.module_rev},{self.CANAddr_AI}')

            # 获取测试信息
            self.isTestCANRunErr = self.checkBox.isChecked()
            self.isTestRunErr = self.checkBox_2.isChecked()

            # self.itemOperation(mTable, 0, 3, 0, '')
            self.itemOperation(mTable, 4, 3, 0, '')
            if self.isTestRunErr:
                self.itemOperation(mTable, 0, 3, 0, '')
                self.itemOperation(mTable, 1, 3, 0, '')
            elif not self.isTestRunErr:
                self.itemOperation(mTable, 0, 0, 0, '')
                self.itemOperation(mTable, 1, 0, 0, '')
            if self.isTestCANRunErr:
                self.itemOperation(mTable, 2, 3, 0, '')
                self.itemOperation(mTable, 3, 3, 0, '')
            elif not self.isTestCANRunErr:
                self.itemOperation(mTable, 2, 0, 0, '')
                self.itemOperation(mTable, 3, 0, 0, '')

            self.inf_param = [mTable, self.module_1, self.module_2, self.testNum]
            self.inf_product = [self.module_type, self.module_pn, self.module_sn, self.module_rev, self.m_Channels]
            self.inf_CANAdrr = [self.CANAddr_DI, self.CANAddr_DO]
            self.inf_additional = [self.baud_rate, self.waiting_time, self.loop_num, self.saveDir]
            self.inf_test = [self.isTestCANRunErr, self.isTestRunErr]
            self.inf_DIDOlist = [self.inf_param, self.inf_product, self.inf_CANAdrr, self.inf_additional, self.inf_test]
            # self.textBrowser_5.clear()
            # self.textBrowser_5.insertPlainText(f'产品信息下发成功，可以开始测试。' + self.HORIZONTAL_LINE)
            # self.pushButton_4.setEnabled(True)
            # self.pushButton_4.setStyleSheet("color: rgb(255, 255, 255);\n"
            #                                 "font: bold 25pt \"宋体\";\n"
            #                                 "background-color: rgb(76, 165, 132);")


        elif self.tabIndex == 1:  # AI界面
            mTable = self.tableWidget_AI
            # # print(f'tabIndex={self.tabIndex}')
            self.module_1 = self.inf['AO1']
            self.module_2 = self.inf['AI2']
            self.testNum = 4  # CAN_RunErr检测 + RunErr检测 + 电流检测 + 电压检测
            self.inf_param = [mTable, self.module_1, self.module_2, self.testNum]
            # 获取产品信息
            self.module_type = self.comboBox_3.currentText()
            self.m_Channels = int(self.module_type[2:4])
            self.AI_Channels = self.m_Channels
            self.module_pn = self.lineEdit_22.text()
            self.module_sn = self.lineEdit_21.text()
            self.module_rev = self.lineEdit_20.text()

            self.inf_product = [self.module_type, self.module_pn, self.module_sn, self.module_rev, self.m_Channels]
            # # print(f'{self.module_type},{self.module_pn},{self.module_sn},{self.module_rev},{self.m_Channels}')
            # 获取CAN地址
            self.CANAddr_AI = int(self.lineEdit_19.text())
            self.CANAddr_AO = int(self.lineEdit_18.text())
            self.CAN1 = self.CANAddr_AO
            self.CAN2 = self.CANAddr_AI
            self.CANAddr_relay = int(self.lineEdit_23.text())
            # self.CANAddr_relay = 0x203
            self.inf_CANAdrr = [self.CANAddr_AI, self.CANAddr_AO, self.CANAddr_relay]
            # 获取附加参数
            self.baud_rate = int(self.lineEdit_16.text())
            self.waiting_time = int(self.lineEdit_17.text())
            self.receive_num = int(self.lineEdit_15.text())
            self.saveDir = self.label_41.text()  # 保存路径
            # # print(f'{self.module_type},{self.module_pn},{self.module_sn},{self.module_rev},{self.CANAddr_AI}')
            self.inf_additional = [self.baud_rate, self.waiting_time, self.receive_num, self.saveDir]
            # 获取标定信息
            if (self.radioButton.isChecked()):
                self.isCalibrate = False
                # # print(f'isCalibrate:{self.isCalibrate}')
                self.isCalibrateVol = False
                self.isCalibrateCur = False
                self.inf_calibrate = [self.isCalibrate, self.isCalibrateVol, self.isCalibrateCur, 0]
            elif (self.radioButton_2.isChecked() or self.radioButton_3.isChecked()):
                self.isCalibrate = True
                # print(f'isCalibrate:{self.isCalibrate}')
                self.isCalibrateVol = self.checkBox_5.isChecked()
                # print(f'isCalibrateVol:{self.isCalibrateVol}')
                self.isCalibrateCur = self.checkBox_6.isChecked()
                # print(f'isCalibrateCur:{self.isCalibrateCur}')
                if self.radioButton_2.isChecked():
                    self.inf_calibrate = [self.isCalibrate, self.isCalibrateVol, self.isCalibrateCur, 1]
                elif self.radioButton_3.isChecked():
                    self.inf_calibrate = [self.isCalibrate, self.isCalibrateVol, self.isCalibrateCur, 2]
                # elif not self.radioButton_2.isChecked() and not self.radioButton_3.isChecked():
                #     self.inf_calibrate = [self.isCalibrate, self.isCalibrateVol, self.isCalibrateCur,0]
            else:
                self.inf_calibrate = [self.isCalibrate, self.isCalibrateVol, self.isCalibrateCur, 0]
            # 获取检测信息
            if (self.radioButton_4.isChecked()):
                self.isTest = False
                # # print(f'isCalibrate:{self.isCalibrate}')
                self.isAITestVol = False
                self.isAITestCur = False
                self.isTestCANRunErr = False
                self.isTestRunErr = False
            elif (self.radioButton_5.isChecked()):
                self.isTest = True
                # # print(f'isTest:{self.isTest}')
                self.isAITestVol = self.checkBox_9.isChecked()
                # # print(f'isAITestVol:{self.isAITestVol}')
                self.isAITestCur = self.checkBox_10.isChecked()
                # # print(f'isAITestCur:{self.isAITestCur}')
                self.isTestCANRunErr = self.checkBox_8.isChecked()
                # # print(f'isTestCANRunErr:{self.isTestCANRunErr}')
                self.isTestRunErr = self.checkBox_7.isChecked()
                # # print(f'isTestRunErr:{self.isTestRunErr}')
            else:
                self.isTest = False
                # # print(f'isCalibrate:{self.isCalibrate}')
                self.isAITestVol = False
                self.isAITestCur = False
                self.isTestCANRunErr = False
                self.isTestRunErr = False
            self.inf_test = [self.isTest, self.isAITestVol, self.isAITestCur, self.isTestCANRunErr, self.isTestRunErr]

            self.itemOperation(mTable, 0, 3, 0, '')
            if self.isTestRunErr:
                self.itemOperation(mTable, 1, 3, 0, '')
                self.itemOperation(mTable, 2, 3, 0, '')
            elif not self.isTestRunErr:
                self.itemOperation(mTable, 1, 0, 0, '')
                self.itemOperation(mTable, 2, 0, 0, '')
            if self.isTestCANRunErr:
                self.itemOperation(mTable, 3, 3, 0, '')
                self.itemOperation(mTable, 4, 3, 0, '')
            elif not self.isTestCANRunErr:
                self.itemOperation(mTable, 3, 0, 0, '')
                self.itemOperation(mTable, 4, 0, 0, '')
            if self.isCalibrateVol:
                self.itemOperation(mTable, 5, 3, 0, '')
            elif not self.isCalibrateVol:
                self.itemOperation(mTable, 5, 0, 0, '')
            if self.isCalibrateCur:
                self.itemOperation(mTable, 6, 3, 0, '')
            elif not self.isCalibrateCur:
                self.itemOperation(mTable, 6, 0, 0, '')
            if self.isAITestVol:
                self.itemOperation(mTable, 7, 3, 0, '')
            elif not self.isAITestVol:
                self.itemOperation(mTable, 7, 0, 0, '')
            if self.isAITestCur:
                self.itemOperation(mTable, 8, 3, 0, '')
            elif not self.isAITestCur:
                self.itemOperation(mTable, 8, 0, 0, '')

            self.inf_AIlist = [self.inf_param, self.inf_product, self.inf_CANAdrr,
                               self.inf_additional, self.inf_calibrate, self.inf_test]

        elif self.tabIndex == 2:  # AO界面
            mTable = self.tableWidget_AO
            # print(f'tabIndex={self.tabIndex}')
            self.module_1 = self.inf['AI1']
            self.module_2 = self.inf['AO2']
            self.testNum = 4  # CAN_RunErr检测 + RunErr检测 + 电流检测 + 电压检测
            self.inf_param = [mTable, self.module_1, self.module_2, self.testNum]
            # 获取产品信息
            self.module_type = self.comboBox_5.currentText()
            self.m_Channels = int(self.module_type[4:])
            self.AO_Channels = self.m_Channels
            self.module_pn = self.lineEdit_47.text()
            self.module_sn = self.lineEdit_46.text()
            self.module_rev = self.lineEdit_45.text()
            # print(f'{self.module_type},{self.module_pn},{self.module_sn},{self.module_rev},{self.m_Channels}')
            self.inf_product = [self.module_type, self.module_pn, self.module_sn, self.module_rev, self.m_Channels]

            # 获取CAN地址
            self.CANAddr_AI = int(self.lineEdit_39.text())
            self.CANAddr_AO = int(self.lineEdit_40.text())
            self.CAN1 = self.CANAddr_AI
            self.CAN2 = self.CANAddr_AO
            self.CANAddr_relay = int(self.lineEdit_41.text())
            # self.CANAddr_relay = 0x203
            self.inf_CANAdrr = [self.CANAddr_AI, self.CANAddr_AO, self.CANAddr_relay]
            # 获取附加参数
            self.baud_rate = int(self.lineEdit_42.text())
            self.waiting_time = int(self.lineEdit_43.text())
            self.receive_num = int(self.lineEdit_44.text())
            self.saveDir = self.label_41.text()  # 保存路径
            # # print(f'{self.module_type},{self.module_pn},{self.module_sn},{self.module_rev},{self.CANAddr_AI}')
            self.inf_additional = [self.baud_rate, self.waiting_time, self.receive_num, self.saveDir]
            # 获取标定信息
            if (self.radioButton_18.isChecked()):
                self.isCalibrate = False
                # # print(f'isCalibrate:{self.isCalibrate}')
                self.isCalibrateVol = False
                self.isCalibrateCur = False
                self.inf_calibrate = [self.isCalibrate, self.isCalibrateVol, self.isCalibrateCur, 0]
            elif (self.radioButton_19.isChecked() or self.radioButton_20.isChecked()):
                self.isCalibrate = True
                # print(f'isCalibrate:{self.isCalibrate}')
                self.isCalibrateVol = self.checkBox_35.isChecked()
                # print(f'isCalibrateVol:{self.isCalibrateVol}')
                self.isCalibrateCur = self.checkBox_36.isChecked()
                # print(f'isCalibrateCur:{self.isCalibrateCur}')
                if self.radioButton_19.isChecked():
                    self.inf_calibrate = [self.isCalibrate, self.isCalibrateVol, self.isCalibrateCur, 1]
                elif self.radioButton_20.isChecked():
                    self.inf_calibrate = [self.isCalibrate, self.isCalibrateVol, self.isCalibrateCur, 2]
            else:
                self.inf_calibrate = [self.isCalibrate, self.isCalibrateVol, self.isCalibrateCur, 0]
            # 获取检测信息
            if (self.radioButton_16.isChecked()):
                self.isTest = False
                # # print(f'isCalibrate:{self.isCalibrate}')
                self.isAOTestVol = False
                self.isAOTestCur = False
                self.isTestCANRunErr = False
                self.isTestRunErr = False
            elif (self.radioButton_17.isChecked()):
                self.isTest = True
                # # print(f'isTest:{self.isTest}')
                self.isAOTestVol = self.checkBox_29.isChecked()
                # # print(f'isAOTestVol:{self.isAOTestVol}')
                self.isAOTestCur = self.checkBox_30.isChecked()
                # # print(f'isAOTestCur:{self.isAOTestCur}')
                self.isTestCANRunErr = self.checkBox_31.isChecked()
                # # print(f'isTestCANRunErr:{self.isTestCANRunErr}')
                self.isTestRunErr = self.checkBox_32.isChecked()
                # # print(f'isTestRunErr:{self.isTestRunErr}')
            else:
                self.isTest = False
                # # print(f'isCalibrate:{self.isCalibrate}')
                self.isAOTestVol = False
                self.isAOTestCur = False
                self.isTestCANRunErr = False
                self.isTestRunErr = False
            self.inf_test = [self.isTest, self.isAOTestVol, self.isAOTestCur, self.isTestCANRunErr, self.isTestRunErr]

            self.itemOperation(mTable, 0, 3, 0, '')
            if self.isTestRunErr:
                self.itemOperation(mTable, 1, 3, 0, '')
                self.itemOperation(mTable, 2, 3, 0, '')
            elif not self.isTestRunErr:
                self.itemOperation(mTable, 1, 0, 0, '')
                self.itemOperation(mTable, 2, 0, 0, '')
            if self.isTestCANRunErr:
                self.itemOperation(mTable, 3, 3, 0, '')
                self.itemOperation(mTable, 4, 3, 0, '')
            elif not self.isTestCANRunErr:
                self.itemOperation(mTable, 3, 0, 0, '')
                self.itemOperation(mTable, 4, 0, 0, '')
            if self.isCalibrateVol:
                self.itemOperation(mTable, 5, 3, 0, '')
            elif not self.isCalibrateVol:
                self.itemOperation(mTable, 5, 0, 0, '')
            if self.isCalibrateCur:
                self.itemOperation(mTable, 6, 3, 0, '')
            elif not self.isCalibrateCur:
                self.itemOperation(mTable, 6, 0, 0, '')
            if self.isAOTestVol:
                self.itemOperation(mTable, 7, 3, 0, '')
            elif not self.isAOTestVol:
                self.itemOperation(mTable, 7, 0, 0, '')
            if self.isAOTestCur:
                self.itemOperation(mTable, 8, 3, 0, '')
            elif not self.isAOTestCur:
                self.itemOperation(mTable, 8, 0, 0, '')

            self.inf_AOlist = [self.inf_param, self.inf_product, self.inf_CANAdrr,
                               self.inf_additional, self.inf_calibrate, self.inf_test]


        elif self.tabIndex == 3:  # CPU界面
            mTable = self.tableWidget_CPU
            # print(f'tabIndex={self.tabIndex}')
            self.module_1 = '工装1（ET1600）'
            self.module_2 = '工装2（QN0016）'
            self.module_3 = '工装3（QN0016）'
            self.module_4 = '工装4（AE0400）'
            self.module_5 = '工装5（AQ0004）'
            self.testNum = 18  # ["外观检测", "型号检查", "SRAM", "FLASH", "拨杆测试", "MFK按键",
            #  "掉电保存", "RTC测试", "FPGA","各指示灯", "本体IN", "本体OUT", "以太网",
            # "RS-232C", "RS-485","右扩CAN", "MAC/三码写入", "MA0202"]

            self.inf_param = [mTable, self.module_1, self.module_2, self.module_3,
                              self.module_4, self.module_5, self.testNum]
            # 获取产品信息
            self.module_type = self.comboBox_20.currentText()
            self.in_Channels = int(self.module_type[5:7])
            self.out_Channels = int(self.module_type[7:9])
            self.module_pn = self.lineEdit_30.text()
            self.module_sn = self.lineEdit_31.text()
            self.module_rev = self.lineEdit_32.text()
            self.module_MAC = self.lineEdit_33.text()
            # print(f'{self.module_type},{self.module_pn},{self.module_sn},{self.module_rev},{self.m_Channels}')
            self.inf_product = [self.module_type, self.module_pn, self.module_sn,
                                self.module_rev, self.module_MAC, self.in_Channels, self.out_Channels]

            # 获取CAN与IP地址
            self.CANAddr1 = int(self.lineEdit_34.text())
            self.CANAddr2 = int(self.lineEdit_35.text())
            self.CANAddr3 = int(self.lineEdit_36.text())
            self.CANAddr4 = int(self.lineEdit_37.text())
            self.CANAddr5 = int(self.lineEdit_38.text())

            # self.IPAddr = '192.168.1.66'
            self.inf_CANIPAdrr = [self.CANAddr1, self.CANAddr2, self.CANAddr3, self.CANAddr4,
                                  self.CANAddr5]
            # 获取串口信息
            self.serialPort_232 = self.comboBox_21.currentText()
            self.serialPort_485 = self.comboBox_22.currentText()
            self.serialPort_typeC = self.comboBox_23.currentText()
            self.saveDir = self.label_41.text()  # 保存路径
            self.serialPort_power = self.comboBox_power.currentText()
            self.serialPort_Op232 = self.comboBox_24.currentText()
            self.serialPort_Op485 = self.comboBox_25.currentText()
            # # print(f'{self.module_type},{self.module_pn},{self.module_sn},{self.module_rev},{self.CANAddr_AI}')
            self.inf_serialPort = [self.serialPort_232, self.serialPort_485,
                                   self.serialPort_typeC, self.saveDir,
                                   self.serialPort_power, self.serialPort_Op232,
                                   self.serialPort_Op485]
            # 获取检测信息

            self.CPU_test_array = [self.checkBox_50, self.checkBox_51, self.checkBox_52, self.checkBox_53,
                                   self.checkBox_55, self.checkBox_56, self.checkBox_57,
                                   self.checkBox_58, self.checkBox_59, self.checkBox_60, self.checkBox_61,
                                   self.checkBox_62, self.checkBox_63, self.checkBox_64, self.checkBox_65,
                                   self.checkBox_66, self.checkBox_54, self.checkBox_67
                                   ]
            self.CPU_testName_array = ["外观检测", "型号检查", "SRAM", "FLASH", "拨杆测试", "MFK按键",
                                       "掉电保存", "RTC测试", "FPGA", "各指示灯", "本体IN", "本体OUT",
                                       "以太网", "RS-232C", "RS-485", "右扩CAN", "MAC/三码写入", "MA0202"]
            self.inf_CPU_test = [False for x in range(len(self.CPU_test_array))]
            for i in range(len(self.CPU_test_array)):
                self.inf_CPU_test[i] = self.CPU_test_array[i].isChecked()
                # print(f' self.CPU_test_array[{i}].isChecked():{ self.CPU_test_array[i].isChecked()}')
            self.inf_test = [self.inf_CPU_test, self.checkBox_68.isChecked()]
            # :param
            # mTable: 进行操作的表格
            # :param
            # row: 进行操作的单元行
            # :param
            # state: 测试状态['本次不检测', '正在检测', '检测完成', '等待检测']
            # :param
            # result: 测试结果
            # col2 = ['', '通过', '未通过']
            # :param
            # operationTime: 测试时间
            for i in range(len(self.CPU_test_array)):
                if self.inf_CPU_test[i]:
                    self.itemOperation(mTable, i, 3, 0, '')
                else:
                    self.itemOperation(mTable, i, 0, 0, '')
            self.inf_CPUlist = [self.inf_param, self.inf_product, self.inf_CANIPAdrr,
                                self.inf_serialPort, self.inf_test, self.current_dir, 'CPU']
        elif self.tabIndex == 4:  # MA0202界面
            # mTable = self.tableWidget_CPU
            self.module_pn = self.lineEdit_MA0202_PN.text()
            self.module_sn = self.lineEdit_MA0202_SN.text()
            self.module_rev = self.lineEdit_MA0202_REV.text()
            # self.module_MAC = self.lineEdit_33.text()
            self.MA0202_isChangePara = self.checkBox_MA0202_para.isChecked()
            self.moduleName_1 = '工装1（AE0400）'
            self.moduleName_2 = '工装2（AQ0004）'
            self.MA0202_CANAddr1 = int(self.lineEdit_MA0202_AE.text())
            self.MA0202_CANAddr2 = int(self.lineEdit_MA0202_AQ.text())
            self.isTestModule = self.radioButton_MA0202.isChecked()
            self.testModuleType = self.radioButton_MA0202.text()
            self.serialPort_typeC = self.comboBox_MA0202_typeC.currentText()

            self.inf_MA0202list = [self.MA0202_isChangePara, self.moduleName_1, self.moduleName_2, self.MA0202_CANAddr1,
                                   self.MA0202_CANAddr2, self.isTestModule, self.testModuleType, self.serialPort_typeC]

        if self.tabIndex == 0 or self.tabIndex == 1 or self.tabIndex == 2:
            # 三码转换
            self.asciiCode_pn = (strToASCII(self.lineEdit_PN.text()))
            self.asciiCode_sn = (strToASCII(self.lineEdit_SN.text()))
            self.asciiCode_rev = (strToASCII(self.lineEdit_REV.text()))

            # 三码写入PLC
            if not write3codeToPLC(self.CAN2, [self.asciiCode_pn, self.asciiCode_sn, self.asciiCode_rev]):
                return False
            else:
                self.showInf('三码写入成功！\n\n')

        self.textBrowser_5.clear()
        self.textBrowser_5.insertPlainText(f'产品信息下发成功，开始测试。' + self.HORIZONTAL_LINE)

        # self.pushButton_pause.setEnabled(True)
        # self.pushButton_pause.setVisible(True)
        #
        # self.pushButton_resume.setEnabled(False)
        # self.pushButton_resume.setVisible(False)
        return True

    def clearList(self, array):
        for i in range(len(array)):
            array[i] = 0x00

    # 自动分配节点
    def configCANAddr(self, list):
        if self.tabIndex == 1 or self.tabIndex == 2:
            list = [list[0], list[1], list[2], list[3]]
        elif self.tabIndex == 3:
            list = [list[0], list[1], list[2]]
        elif self.tabIndex == 4:
            list = [list[0], list[1]]
        else:
            list = [list[0], list[1]]
        for a in list:
            self.m_transmitData = [0xac, 0x00, 0x00, 0x00, 0x00, 0x00, 0x00, 0x00]
            self.m_transmitData[2] = a
            boola = CAN_option.transmitCAN(0x0, self.m_transmitData, 1)[0]
            time.sleep(0.01)
            if not boola:
                self.showInf(f'节点{a}分配失败' + self.HORIZONTAL_LINE)
                return False
        self.showInf(f'所有节点分配成功' + self.HORIZONTAL_LINE)
        return True

    def isModulesOnline(self, can_list, module_list):
        # 检测设备心跳
        for i in range(len(can_list)):
            if not self.check_heartbeat(can_list[i], module_list[i], self.waiting_time):
                return False
        return True
        # if self.check_heartbeat(self.CAN1, self.module_1, self.waiting_time) == False:
        #     return False
        # if self.check_heartbeat(self.CAN2, self.module_2, self.waiting_time) == False:
        #     return False
        # if self.tabIndex !=0:
        #     if self.check_heartbeat(self.CANAddr_relay, '继电器QR0016#1', self.waiting_time) == False:
        #         return False
        #     if self.check_heartbeat(self.CANAddr_relay+1, '继电器QR0016#2', self.waiting_time) == False:
        #         return False

    def check_heartbeat(self, can_addr, inf, max_waiting):
        can_id = 0x700 + can_addr
        bool_receive, self.m_can_obj = CAN_option.receiveCANbyID(can_id, max_waiting, 1)

        if bool_receive == False:
            self.showInf(f'错误：未发现{inf}' + self.HORIZONTAL_LINE)
            return False
        else:
            self.showInf(f'发现{inf}：收到心跳帧：{hex(self.m_can_obj.ID)}\n\n')

        return True
        # if inf == '继电器':
        #     bool_receive, self.m_can_obj = CAN_option.receiveCANbyID(0x700 + can_addr , max_waiting,0)
        #     # print(self.m_can_obj.Data)
        #     if bool_receive == False:
        #         self.showInf(f'错误：未发现{inf}' + self.HORIZONTAL_LINE)
        #         # self.isPause()
        #         # if not self.isStop():
        #         #     return
        #         return False
        #
        #     self.showInf(f'发现{inf}：收到心跳帧：{hex(self.m_can_obj.ID)}\n\n')

    def getSerialInf(self, num: int, type):
        # self.textBrowser_5.clear()
        if type == "CPU":
            # 清空串口选项
            for i in range(self.comboBox_21.count()):
                self.comboBox_21.removeItem(i)
                self.comboBox_22.removeItem(i)
                self.comboBox_23.removeItem(i)
                self.comboBox_24.removeItem(i)
                self.comboBox_25.removeItem(i)
                self.comboBox_power.removeItem(i)
            # 获取电脑上可用的串口列表
            ports = serial.tools.list_ports.comports()

            # 遍历串口列表并打印串口信息
            for i in range(len(ports)):
                self.comboBox_21.addItem("")
                self.comboBox_22.addItem("")
                self.comboBox_23.addItem("")
                self.comboBox_24.addItem("")
                self.comboBox_25.addItem("")
                self.comboBox_power.addItem("")
                self.comboBox_21.setItemText(i, ports[i].device)
                self.comboBox_22.setItemText(i, ports[i].device)
                self.comboBox_23.setItemText(i, ports[i].device)
                self.comboBox_24.setItemText(i, ports[i].device)
                self.comboBox_25.setItemText(i, ports[i].device)
                self.comboBox_power.setItemText(i, ports[i].device)
                if num != 0:
                    if i == 0:
                        self.showInf(f'可用串口：\n')
                    self.showInf(f'({i + 1}){ports[i].description}\n')
                # self.showInf(f'({i+1}){ports[i].device}, {ports[i].name}, {ports[i].description}\n')

            self.comboBox_21.removeItem(len(ports))
            self.comboBox_22.removeItem(len(ports))
            self.comboBox_23.removeItem(len(ports))
            self.comboBox_24.removeItem(len(ports))
            self.comboBox_25.removeItem(len(ports))
            self.comboBox_power.removeItem(len(ports))
        elif type == "MA0202":
            # 清空串口选项
            for i in range(self.comboBox_MA0202_typeC.count()):
                self.comboBox_MA0202_typeC.removeItem(i)
            # 获取电脑上可用的串口列表
            ports = serial.tools.list_ports.comports()

            # 遍历串口列表并打印串口信息
            for i in range(len(ports)):
                self.comboBox_MA0202_typeC.addItem("")
                self.comboBox_MA0202_typeC.setItemText(i, ports[i].device)

                if num != 0:
                    if i == 0:
                        self.showInf(f'可用串口：\n')
                    self.showInf(f'({i + 1}){ports[i].description}\n')
                # self.showInf(f'({i+1}){ports[i].device}, {ports[i].name}, {ports[i].description}\n')

            self.comboBox_MA0202_typeC.removeItem(len(ports))

    def CANFail(self):
        self.endOfTest()
        self.initPara()
        # self.textBrowser_5.clear()
        self.errorInf = ''
        self.errorNum = 0
        self.isAIPassTest = True
        self.isAIVolPass = True
        self.isAICurPass = True

        self.isAOPassTest = True
        self.isAOVolPass = True
        self.isAOCurPass = True

        self.isLEDRunOK = True
        self.isLEDErrOK = True
        self.CAN_runLED = True
        self.CAN_errorLED = True
        self.reInputPNSNREV()

    def saveExcel(self, saveList):
        saveList[0].save(str(self.label_41.text()) + saveList[1])
        # book.save(self.label_41.text() + saveDir)

    # def printResult(self,list):
    #     try:
    #         self.generateLabel(list)
    #         content = list[1]
    #         file_name = f'{self.label_41.text()}{list[0]}_label.docx'
    #         # if list[1] == 'FAIL':
    #         #     content += f'\n\n{list[2]}'
    #         #
    #         # with open(file_name, "w") as f:
    #         #     for line in content.splitlines():
    #         #         f.write(line + "\n")
    #
    #         win32api.ShellExecute(
    #             0,
    #             "print",
    #             file_name,
    #             #
    #             # If this is None, the default printer will
    #             # be used anyway.
    #             #
    #             '/d:"%s"' % win32print.GetDefaultPrinter(),
    #             ".",
    #             0
    #         )
    #     except Exception as e:
    #         self.showInf(f"printResultError:{e}+{self.HORIZONTAL_LINE}")
    #         traceback.print_exc()

    def generateLabel(self, list):
        from docx import Document

        # docx.shared 用于设置大小（图片等）

        from docx.shared import Cm, Pt
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        from docx.document import Document as Doc
        try:
            # 创建代表Word文档的Doc对象

            document = Document()
            default_section = document.sections[0]
            # 默认宽度和高度
            # # print(default_section.page_width.cm)  # 21.59
            # # print(default_section.page_height.cm)  # 27.94
            # 可直接修改宽度和高度，即纸张大小改为自定义
            default_section.page_width = Cm(8)
            default_section.page_height = Cm(12)

            # 修改页边距
            default_section.top_margin = Cm(0.5)
            default_section.right_margin = Cm(0.5)
            default_section.bottom_margin = Cm(0.5)
            default_section.left_margin = Cm(0.5)

            # 添加图片（注意路径和图片必须要存在）
            document.add_picture(self.current_dir + '/logo.png', width=Cm(6.1))

            # # 添加带样式的段落
            p = document.add_paragraph('')
            p.paragraph_format.line_spacing = 1
            PLC = p.add_run('PLC I/O')
            PLC.font.name = 'Stencil'
            PLC.font.size = Pt(24)
            PLC.underline = False

            s = p.add_run('s fac_Test')
            s.font.name = 'Stencil'
            s.font.size = Pt(16)
            s.underline = False

            pSN = document.add_paragraph('')
            pSN.paragraph_format.line_spacing = 1
            SN = pSN.add_run(f'UUT SN: {self.module_sn}')
            SN.font.name = '等线'
            SN.font.size = Pt(10.5)

            pPN = document.add_paragraph('')
            pPN.paragraph_format.line_spacing = 1
            PN = pPN.add_run(f'UUT PN: {self.module_pn}')
            PN.font.name = '等线'
            PN.font.size = Pt(10.5)

            pREV = document.add_paragraph('')
            pREV.paragraph_format.line_spacing = 1
            REV = pREV.add_run(f'UUT REV: {self.module_rev}')
            REV.font.name = '等线'
            REV.font.size = Pt(10.5)

            pTime = document.add_paragraph('')
            pTime.paragraph_format.line_spacing = 1
            Time = pTime.add_run(f'Time Duration: {round(time.time() - self.allStart_time, 1)} 秒')
            Time.font.name = '等线'
            Time.font.size = Pt(12)
            Time.bold = True
            if list[1] == 'FAIL':
                # 添加图片（注意路径和图片必须要存在）
                document.add_picture(self.current_dir + '/fail.png', width=Cm(5.5))
                pFail = document.add_paragraph('')
                pFail.paragraph_format.line_spacing = 1
                Fail = pFail.add_run('FAILED ITEMS：')
                Fail.font.name = '等线'
                Fail.font.size = Pt(12)
                Fail.bold = True

                pFail_inf = document.add_paragraph('')
                pFail_inf.paragraph_format.line_spacing = 1
                fail_inf = pFail_inf.add_run(f'{list[2][1:]}')
                fail_inf.font.name = '等线'
                fail_inf.font.size = Pt(10)
            elif list[1] == 'PASS':
                # 添加图片（注意路径和图片必须要存在）
                document.add_picture(self.current_dir + '/pass.png', width=Cm(5.5))
            document.paragraphs[6].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # 保存文档
            for p in document.paragraphs:
                p.paragraph_format.line_spacing = 1
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
            document.save(self.label_41.text() + f'{list[0]}_label.docx')
        except Exception as e:
            self.showInf(f"generateLabelError:{e}+{self.HORIZONTAL_LINE}")
            # 捕获异常并输出详细的错误信息
            traceback.print_exc()

    @abstractmethod
    def calibrateAO(self):
        raise NotImplementedError()

    @abstractmethod
    def testAO(self):
        raise NotImplementedError()

    @abstractmethod
    def calibrateAI(self):
        raise NotImplementedError()

    @abstractmethod
    def testAI(self):
        raise NotImplementedError()

    @abstractmethod
    def testDI(self):
        raise NotImplementedError()

    @abstractmethod
    def testDO(self):
        raise NotImplementedError()

    def showInf(self, inf):
        global isMainRunning
        if isMainRunning:
            self.textBrowser_5.append(inf)
            if inf[:8] == '本轮测试总时间：':
                self.move_to_end()
            QApplication.processEvents()
            time.sleep(0.1)
        else:
            self.textBrowser_5.append('等待进程停止。\n')
            QApplication.processEvents()
            time.sleep(0.1)

        # self.isPause()
        # if self.work_thread.stopFlag.isSet():
        #     return

    def move_to_end(self):
        self.textBrowser_5.moveCursor(self.textBrowser_5.textCursor().End)

    # def isPause(self):
    #     # # print(f'self.work_thread.pauseFlag:{self.work_thread.pauseFlag.isSet()}')
    #     if self.work_thread.pauseFlag.isSet():
    #         while True:
    #             # # print(f'while:{self.pause_num}')
    #             if self.pause_num == 1 and not self.work_thread.stopFlag.isSet():
    #                 self.pause_num += 1
    #                 self.showInf(self.HORIZONTAL_LINE + '暂停中…………'+self.HORIZONTAL_LINE)
    #
    #             QApplication.processEvents()
    #             # # print('暂停中…………')
    #             if not self.work_thread.pauseFlag.isSet():
    #                 self.pause_num = 1
    #                 break

    def isStop(self):
        if self.work_thread.stopFlag.isSet():
            return False
        return True

    def endOfTest(self):
        self.pushButton_pause.setEnabled(False)
        self.pushButton_pause.setVisible(False)

        self.pushButton_resume.setEnabled(False)
        self.pushButton_resume.setVisible(False)

        self.pushButton_4.setStyleSheet(self.topButton_qss['off'])
        self.pushButton_4.setEnabled(False)
        self.pushButton_4.setVisible(True)

        self.pushButton_3.setStyleSheet(self.topButton_qss['off'])
        self.pushButton_3.setEnabled(False)
        self.reInputPNSNREV()
        # self.pushButton_9.setEnabled(False)
        # self.pushButton_9.setStyleSheet(self.topButton_qss['off'])

    def AI_itemOperation(self, list):
        '''
        :param list = [row,state,result,operationTime]
        :param row: 进行操作的单元行
        :param state: 测试状态 ['本次不检测','正在检测','检测完成','等待检测']
        :param result: 测试结果 col2 = ['','通过','未通过']
        :param operationTime: 测试时间
        :return:
        '''
        mTable = self.tableWidget_AI
        col1 = ['本次不检测', '正在检测', '检测完成', '等待检测']
        col2 = ['', '通过', '未通过']
        font = QtGui.QColor(0, 0, 0)
        if list[1] == 0:
            color = QtGui.QColor(255, 255, 255)
            font = QtGui.QColor(197, 197, 197)
        elif list[1] == 1:
            color = QtGui.QColor(255, 255, 0)
        elif list[1] == 2 and list[2] == 1:
            color = QtGui.QColor(0, 255, 0)
        elif list[1] == 2 and list[2] == 2:
            color = QtGui.QColor(255, 0, 0)
            font = QtGui.QColor(255, 255, 255)
        elif list[1] == 3:
            color = QtGui.QColor(197, 197, 197)

        for i in range(4):
            item = mTable.item(list[0], i)
            item.setBackground(color)
            item.setForeground(font)
            item.setTextAlignment(QtCore.Qt.AlignCenter)
            if i == 1:
                item.setText(col1[list[1]])
            if i == 2:
                item.setText(col2[list[2]])
            if i == 3 and list[1] == 2:
                item.setText(f'{list[3]}s')
            if (i == 3 and list[1] == 0) or (i == 3 and list[1] == 3):
                item.setText(f'{list[3]}')
        QApplication.processEvents()

    def AO_itemOperation(self, list):
        '''
        :param list = [row,state,result,operationTime]
        :param row: 进行操作的单元行
        :param state: 测试状态 ['本次不检测','正在检测','检测完成','等待检测']
        :param result: 测试结果 col2 = ['','通过','未通过']
        :param operationTime: 测试时间
        :return:
        '''
        mTable = self.tableWidget_AO
        col1 = ['本次不检测', '正在检测', '检测完成', '等待检测']
        col2 = ['', '通过', '未通过']
        font = QtGui.QColor(0, 0, 0)
        if list[1] == 0:
            color = QtGui.QColor(255, 255, 255)
            font = QtGui.QColor(197, 197, 197)
        elif list[1] == 1:
            color = QtGui.QColor(255, 255, 0)
        elif list[1] == 2 and list[2] == 1:
            color = QtGui.QColor(0, 255, 0)
        elif list[1] == 2 and list[2] == 2:
            color = QtGui.QColor(255, 0, 0)
            font = QtGui.QColor(255, 255, 255)
        elif list[1] == 3:
            color = QtGui.QColor(197, 197, 197)

        for i in range(4):
            item = mTable.item(list[0], i)
            item.setBackground(color)
            item.setForeground(font)
            item.setTextAlignment(QtCore.Qt.AlignCenter)
            if i == 1:
                item.setText(col1[list[1]])
            if i == 2:
                item.setText(col2[list[2]])
            if i == 3 and list[1] == 2:
                item.setText(f'{list[3]}s')
            if (i == 3 and list[1] == 0) or (i == 3 and list[1] == 3):
                item.setText(f'{list[3]}')
        QApplication.processEvents()

    def DIDO_itemOperation(self, list):
        '''
        :param list = [row,state,result,operationTime]
        :param row: 进行操作的单元行
        :param state: 测试状态 ['本次不检测','正在检测','检测完成','等待检测']
        :param result: 测试结果 col2 = ['','通过','未通过']
        :param operationTime: 测试时间
        :return:
        '''
        mTable = self.tableWidget_DIDO
        col1 = ['本次不检测', '正在检测', '检测完成', '等待检测']
        col2 = ['', '通过', '未通过']
        font = QtGui.QColor(0, 0, 0)
        if list[1] == 0:
            color = QtGui.QColor(255, 255, 255)
            font = QtGui.QColor(197, 197, 197)
        elif list[1] == 1:
            color = QtGui.QColor(255, 255, 0)
        elif list[1] == 2 and list[2] == 1:
            color = QtGui.QColor(0, 255, 0)
        elif list[1] == 2 and list[2] == 2:
            color = QtGui.QColor(255, 0, 0)
            font = QtGui.QColor(255, 255, 255)
        elif list[1] == 3:
            color = QtGui.QColor(197, 197, 197)

        for i in range(4):
            item = mTable.item(list[0], i)
            item.setBackground(color)
            item.setForeground(font)
            item.setTextAlignment(QtCore.Qt.AlignCenter)
            if i == 1:
                item.setText(col1[list[1]])
            if i == 2:
                item.setText(col2[list[2]])
            if i == 3 and list[1] == 2:
                item.setText(f'{list[3]}s')
            if (i == 3 and list[1] == 0) or (i == 3 and list[1] == 3):
                item.setText(f'{list[3]}')
        QApplication.processEvents()

    def CPU_itemOperation(self, list):
        '''
        :param list = [row,state,result,operationTime]
        :param row: 进行操作的单元行
        :param state: 测试状态 ['本次不检测','正在检测','检测完成','等待检测']
        :param result: 测试结果 col2 = ['','通过','未通过']
        :param operationTime: 测试时间
        :return:
        '''
        mTable = self.tableWidget_CPU
        col1 = ['本次不检测', '正在检测', '检测完成', '等待检测']
        col2 = ['', '通过', '未通过']
        font = QtGui.QColor(0, 0, 0)
        if list[1] == 0:
            color = QtGui.QColor(255, 255, 255)
            font = QtGui.QColor(197, 197, 197)
        elif list[1] == 1:
            color = QtGui.QColor(255, 255, 0)
        elif list[1] == 2 and list[2] == 1:
            color = QtGui.QColor(0, 255, 0)
        elif list[1] == 2 and list[2] == 2:
            color = QtGui.QColor(255, 0, 0)
            font = QtGui.QColor(255, 255, 255)
        elif list[1] == 3:
            color = QtGui.QColor(197, 197, 197)

        for i in range(4):
            item = mTable.item(list[0], i)
            item.setBackground(color)
            item.setForeground(font)
            item.setTextAlignment(QtCore.Qt.AlignCenter)
            if i == 1:
                item.setText(col1[list[1]])
            if i == 2:
                item.setText(col2[list[2]])
            if i == 3 and list[1] == 2:
                item.setText(f'{list[3]}s')
            if (i == 3 and list[1] == 0) or (i == 3 and list[1] == 3):
                item.setText(f'{list[3]}')
        QApplication.processEvents()

    def itemOperation(self, mTable, row, state, result, operationTime):
        '''

        :param mTable: 进行操作的表格
        :param row: 进行操作的单元行
        :param state: 测试状态 ['本次不检测','正在检测','检测完成','等待检测']
        :param result: 测试结果 col2 = ['','通过','未通过']
        :param operationTime: 测试时间
        :return:
        '''
        col1 = ['本次不检测', '正在检测', '检测完成', '等待检测']
        col2 = ['', '通过', '未通过']
        font = QtGui.QColor(0, 0, 0)
        if state == 0:
            color = QtGui.QColor(255, 255, 255)
            font = QtGui.QColor(197, 197, 197)
        elif state == 1:
            color = QtGui.QColor(255, 255, 0)
        elif state == 2 and result == 1:
            color = QtGui.QColor(0, 255, 0)
        elif state == 2 and result == 2:
            color = QtGui.QColor(255, 0, 0)
            font = QtGui.QColor(255, 255, 255)
        elif state == 3:
            color = QtGui.QColor(197, 197, 197)

        for i in range(4):
            item = mTable.item(row, i)
            item.setBackground(color)
            item.setForeground(font)
            item.setTextAlignment(QtCore.Qt.AlignCenter)
            if i == 1:
                item.setText(col1[state])
            if i == 2:
                item.setText(col2[result])
            if i == 3 and state == 2:
                item.setText(f'{operationTime}s')
            if (i == 3 and state == 0) or (i == 3 and state == 3):
                item.setText(f'{operationTime}')
        QApplication.processEvents()

    def workerChange(self,inf:list):
        # 弹出文本输入对话框
        if inf[0]=='admin':
            if self.checkBox_admin.isChecked():
                text, ok = QInputDialog.getText(None, inf[1], inf[2],echo=QLineEdit.Password)
            elif not self.checkBox_admin.isChecked():
                #关闭所有参数修改权限
                for g in self.group_list:
                    g.setEnabled(False)
        else:
            text, ok = QInputDialog.getText(None, inf[1], inf[2])
        if inf[0]=='start':
            if len(text) != 0:
                self.label_33.setText(text)
            else:
                self.killUi()
        else:
            if inf[0]=='worker':
                if ok:
                    self.label_33.setText(text)
                    self.label_userName.setText(f'<h2>{text}<h2>')
                    self.checkBox_admin.setChecked(False)
            elif inf[0]=='admin':
                if self.checkBox_admin.isChecked():
                    if ok:
                        if text == 'admin123':
                            # if self.checkBox_admin.isChecked():
                            for g in self.group_list:
                                g.setEnabled(True)
                        else:
                            self.showMessageBox_oneButton(['警告','管理员密码输入错误！'])
                            self.checkBox_admin.setChecked(False)

            # elif inf[0]=='start':
            #     if not ok:
            #         self.killUi()

    def userLogin(self):
        global isMainRunning

        if isMainRunning:
            msg_box = QMessageBox()
            msg_box.setStandardButtons(QMessageBox.Ok)
            user_widget = QWidget()

            layout_user = QHBoxLayout()
            label_user = QLabel('<h2>当前员工：<h2>')
            self.label_userName.setText(f'<h2>{self.label_33.text()}<h2>')

            layout_user.addWidget(label_user)
            layout_user.addWidget(self.label_userName)

            self.checkBox_admin.setChecked(self.adminState)

            pushButton_login = QPushButton('更换员工账号')

            layout = QVBoxLayout()
            layout.addLayout(layout_user)
            layout.addWidget(self.checkBox_admin)
            layout.addWidget(pushButton_login)

            user_widget.setLayout(layout)

            # 将QWidget对象添加到QMessageBox中
            msg_box.layout().addWidget(user_widget)
            # msg_box.setText(f'<center><h2>当前员工：{self.label_33.text()}</h2></center>')


            msg_box.setWindowTitle('员工登录')
            # 设置样式表
            msg_box.setStyleSheet('QLabel{font-size: 18px;}')
            msg_box.button(QMessageBox.Ok).setHidden(True)
            #登录管理员账号
            self.checkBox_admin.toggled.connect(lambda: self.workerChange(['admin','管理员登录', '请输入管理员密码：']))
            self.checkBox_admin.toggled.connect(self.saveAdminState)
            #更换员工账号
            pushButton_login.clicked.connect(lambda: self.workerChange(['worker','测试人员登录', '请输入员工工号：']))

            # 显示消息框
            reply = msg_box.exec_()
            # # 将弹窗结果放入队列
            self.result_queue.put([reply])

    def saveAdminState(self):
        self.adminState = self.checkBox_admin.isChecked()

    def option_pushButton13(self):
        self.lineEdit_SN.setText('S1223-001083')
        self.lineEdit_REV.setText('06')

    def initPara(self):
        # 生成表格的相关标识
        # DI测试是否通过
        self.isDIPassTest = True
        # DO测试是否通过
        self.isDOPassTest = True
        # DO通道数据初始化
        self.DO_channelData = 0
        # DI通道数据初始化
        self.DI_channelData = 0
        # DO通道错误数据记录
        self.DODataCheck = [True for i in range(32)]
        # DI通道错误数据记录
        self.DIDataCheck = [True for i in range(32)]

        # 是否标定
        self.isCalibrate = False
        # 是否标定电压
        self.isCalibrateVol = False
        # 是否标定电流
        self.isCalibrateCur = False

        # 是否检测
        self.isCalibrate = False
        # 是否检测AI电压
        self.isAITestVol = False
        # 是否检测AI电流
        self.isAITestCur = False
        # 是否检测AO电压
        self.isAOTestVol = False
        # 是否检测AO电流
        self.isAOTestCur = False
        # 是否检测CAN_Run+CAN_Error
        self.isTestCANRunErr = False
        # 是否检测Run+Error
        self.isTestRunErr = False
        # AI测试是否通过
        self.isAIPassTest = True
        self.isAIVolPass = True
        self.isAICurPass = True
        # AO测试是否通过
        self.isAOPassTest = True
        self.isAOVolPass = True
        self.isAOCurPass = True

        self.isLEDRunOK = True
        self.isLEDErrOK = True
        self.isLEDPass = True

        # 测试总数
        self.testNum = 0
        # 测试类型
        self.testType = {}
        # 模块信息
        # inf = {'AO2': '待检AO模块', 'AI1': '配套AI模块', 'AI2': '待检AI模块', 'AO1': '配套AO模块',
        #        'DO2': '待检DO模块', 'DI1': '配套DI模块', 'DI2': '待检DI模块', 'DO1': '配套DO模块'}
        self.module_1 = ''
        self.module_2 = ''
        # 通道数
        self.m_Channels = 0
        self.AI_Channels = 0
        self.AO_Channels = 0
        # 发送的数据
        self.ubyte_array_transmit = c_ubyte * 8
        self.m_transmitData = self.ubyte_array_transmit(0x00, 0x00, 0x00, 0x00, 0x00, 0x00, 0x00, 0x00)
        # 接收的数据
        self.ubyte_array_receive = c_ubyte * 8
        self.m_receiveData = self.ubyte_array_receive(0x00, 0x00, 0x00, 0x00, 0x00, 0x00, 0x00, 0x00)
        self.CAN_errorLED = True
        self.CAN_runLED = True
        self.errorLED = True
        self.runLED = True
        self.errorNum = 0
        self.errorInf = ''
        self.isAllScreen = False
        self.volReceValue = [0, 0, 0, 0, 0]
        self.curReceValue = [0, 0, 0, 0, 0]
        self.volPrecision = [0, 0, 0, 0, 0]
        self.curPrecision = [0, 0, 0, 0, 0]
        self.chPrecision = [0, 0, 0, 0]

        self.subRunFlag = True


def CAN_init(CAN_channel: list):
    CAN_option.close(CAN_option.VCI_USB_CAN_2, CAN_option.DEV_INDEX)
    time.sleep(0.1)
    QApplication.processEvents()
    for channel in CAN_channel:
        if not CAN_option.connect(CAN_option.VCI_USB_CAN_2, CAN_option.DEV_INDEX, channel):
            # self.showMessageBox(['CAN设备存在问题', 'CAN设备开启失败，请检查CAN设备！'])
            # self.CANFail()
            return [False, ['CAN设备存在问题', f'CAN通道{channel}开启失败，请检查CAN设备！']]

        if not CAN_option.init(CAN_option.VCI_USB_CAN_2, CAN_option.DEV_INDEX, channel, init_config):
            # self.showMessageBox(['CAN设备存在问题', 'CAN通道初始化失败，请检查CAN设备！'])
            # self.CANFail()
            return [False, ['CAN设备存在问题', f'CAN通道{channel}初始化失败，请检查CAN设备！']]

        if not CAN_option.start(CAN_option.VCI_USB_CAN_2, CAN_option.DEV_INDEX, channel):
            # self.showMessageBox(['CAN设备存在问题','CAN通道打开失败，请检查CAN设备！'])
            # self.CANFail()
            return [False, ['CAN设备存在问题', f'CAN通道{channel}打开失败，请检查CAN设备！']]
    return [True, ['', '']]


def mainThreadRunning():
    global isMainRunning
    if not isMainRunning:
        return False
    return True


def canInit_thread():
    q.put(CAN_init([1]))
    event_canInit.set()


def strToASCII(str):
    ascii_list = []
    for char in str:
        ascii_code = int(ord(char))
        ascii_list.append(ascii_code)

    # print(ascii_list)
    return ascii_list


def ASCIIToHex(list):
    hex_list = [0x00 for x in range(len(list))]
    # print(f'hex_list = {hex_list}')
    num = 0
    for n in list:
        # print(type(hex_list[num]))
        hex_list[num] = hex(n)
        # print(type(hex_list[num]))
        num += 1
    # print(f'hex_list = {hex_list}')
    return hex_list


def write3codeToPLC(addr, code_list):
    ubyte_transmit = c_ubyte * 8
    transmit_inf = ubyte_transmit(0x00, 0x00, 0x00, 0x00, 0x00, 0x00, 0x00, 0x00)
    # 写入PN
    code_pn = code_list[0]
    # print(f'code_pn = {code_pn}')
    num1 = int(len(code_pn) / 4)
    num2 = len(code_pn) % 4
    if num1 != 0:
        for i in range(num1):
            transmit_inf = ubyte_transmit(0x00, 0x00, 0x00, 0x00, 0x00, 0x00, 0x00, 0x00)
            transmit_inf = [0x23, 0xf8, 0x5f, 0x00 + (i + 1), 0x00 + code_pn[i * 4 + 0], 0x00 + code_pn[i * 4 + 1],
                            0x00 + code_pn[i * 4 + 2], 0x00 + code_pn[i * 4 + 3]]
            # transmit_inf = [0x23, 0xf7, 0x5f, 0x01, 0x53, 0x31, 0x32, 0x32]
            bool_pn, rece_pn = CAN_option.transmitCAN(0x602, transmit_inf, 0)
            if not bool_pn:
                return False
    if num2 != 0:
        transmit_inf = ubyte_transmit(0x00, 0x00, 0x00, 0x00, 0x00, 0x00, 0x00, 0x00)
        if num2 == 1:
            transmit_inf[0] = 0x2f
            transmit_inf[4] = code_pn[num1 * 4 + 0]
        elif num2 == 2:
            transmit_inf[0] = 0x2b
            transmit_inf[4] = 0x00 + code_pn[num1 * 4 + 0]
            transmit_inf[5] = 0x00 + code_pn[num1 * 4 + 1]
        elif num2 == 3:
            transmit_inf[0] = 0x27
            transmit_inf[4] = 0x00 + code_pn[num1 * 4 + 0]
            transmit_inf[5] = 0x00 + code_pn[num1 * 4 + 1]
            transmit_inf[6] = 0x00 + code_pn[num1 * 4 + 2]

        transmit_inf[1] = 0xf8
        transmit_inf[2] = 0x5f
        transmit_inf[3] = num1 + 1
        bool_pn, rece_pn = CAN_option.transmitCAN(0x602, transmit_inf, 0)
        if not bool_pn:
            return False

    # 写入SN
    code_sn = code_list[1]
    # print('code_sn = ',code_sn)
    num1 = int(len(code_sn) / 4)
    num2 = len(code_sn) % 4
    if num1 != 0:
        for i in range(num1):
            transmit_inf = ubyte_transmit(0x00, 0x00, 0x00, 0x00, 0x00, 0x00, 0x00, 0x00)
            transmit_inf = [0x23, 0xf7, 0x5f, 0x00 + (i + 1), 0x00 + code_sn[i * 4 + 0], 0x00 + code_sn[i * 4 + 1],
                            0x00 + code_sn[i * 4 + 2], 0x00 + code_sn[i * 4 + 3]]
            # transmit_inf = [0x23, 0xf7, 0x5f, 0x01, 0x53, 0x31,0x32,0x32]
            bool_sn, rece_sn = CAN_option.transmitCAN(0x602, transmit_inf, 0)
            if not bool_sn:
                return False
    if num2 != 0:
        transmit_inf = ubyte_transmit(0x00, 0x00, 0x00, 0x00, 0x00, 0x00, 0x00, 0x00)
        if num2 == 1:
            transmit_inf[0] = 0x2f
            transmit_inf[4] = 0x00 + code_sn[num1 * 4 + 0]
        elif num2 == 2:
            transmit_inf[0] = 0x2b
            transmit_inf[4] = 0x00 + code_sn[num1 * 4 + 0]
            transmit_inf[5] = 0x00 + code_sn[num1 * 4 + 1]
        elif num2 == 3:
            transmit_inf[0] = 0x27
            transmit_inf[4] = 0x00 + code_sn[num1 * 4 + 0]
            transmit_inf[5] = 0x00 + code_sn[num1 * 4 + 1]
            transmit_inf[6] = 0x00 + code_sn[num1 * 4 + 2]

        transmit_inf[1] = 0xf7
        transmit_inf[2] = 0x5f
        transmit_inf[3] = 0x00 + (num1 + 1)
        bool_sn, rece_sn = CAN_option.transmitCAN(0x602, transmit_inf, 0)
        if not bool_sn:
            return False
    # print('code_sn = ', code_sn)
    # 写入REV
    code_rev = code_list[2]
    num1 = int(len(code_rev) / 4)
    num2 = len(code_rev) % 4
    if num1 != 0:
        for i in range(num1):

            transmit_inf = ubyte_transmit(0x00, 0x00, 0x00, 0x00, 0x00, 0x00, 0x00, 0x00)
            transmit_inf = [0x23, 0xf9, 0x5f, 0x00 + (i + 1), 0x00 + code_rev[i * 4 + 0], 0x00 + code_rev[i * 4 + 1],
                            0x00 + code_rev[i * 4 + 2], 0x00 + code_rev[i * 4 + 3]]
            bool_rev, rece_rev = CAN_option.transmitCAN(0x600 + addr, transmit_inf, 0)
            if not bool_rev:
                return False
    if num2 != 0:
        transmit_inf = ubyte_transmit(0x00, 0x00, 0x00, 0x00, 0x00, 0x00, 0x00, 0x00)
        if num2 == 1:
            transmit_inf[0] = 0x2f
            transmit_inf[4] = 0x00 + code_rev[num1 * 4 + 0]
        elif num2 == 2:
            transmit_inf[0] = 0x2b
            transmit_inf[4] = 0x00 + code_rev[num1 * 4 + 0]
            transmit_inf[5] = 0x00 + code_rev[num1 * 4 + 1]
        elif num2 == 3:
            transmit_inf[0] = 0x27
            transmit_inf[4] = 0x00 + code_rev[num1 * 4 + 0]
            transmit_inf[5] = 0x00 + code_rev[num1 * 4 + 1]
            transmit_inf[6] = 0x00 + code_rev[num1 * 4 + 2]

        transmit_inf[1] = 0xf9
        transmit_inf[2] = 0x5f
        transmit_inf[3] = 0x00 + (num1 + 1)
        bool_rev, rece_rev = CAN_option.transmitCAN(0x600 + addr, transmit_inf, 0)
        if not bool_rev:
            return False

    return True


def get3codeFromPLC(addr):
    ubyte_transmit = c_ubyte * 8
    transmit_inf = ubyte_transmit(0x00, 0x00, 0x00, 0x00, 0x00, 0x00, 0x00, 0x00)

    # 读设备PN
    pn_list = []
    flag_pn = True
    for i in range(6):
        transmit_inf[0] = 0x40
        transmit_inf[1] = 0xf8
        transmit_inf[2] = 0x5f
        transmit_inf[3] = 0x00 + i + 1
        bool_pn, transmit_pn = CAN_option.transmitCAN(0x602, transmit_inf, 0)
        if not bool_pn:
            return False

        while True:
            bool_pn, m_can_obj = CAN_option.receiveCANbyID(0x582, 2, 0)
            if bool_pn:
                break
            if bool_pn == 'stopReceive':
                return 'stopReceive', m_can_obj.Data
            elif not bool_pn:
                return False, m_can_obj.Data

        for j in range(4):
            if m_can_obj.Data[j + 4] != 0x00:
                pn_list.append(int(m_can_obj.Data[j + 4]))
            else:
                flag_pn = False
                break

        if not flag_pn:
            break

    # print('pn_list = ',pn_list)
    pnCode = ''
    for asc in pn_list:
        pnCode += chr(asc)
    # print('pnCode = ', pnCode)

    # 读设备SN
    sn_list = []
    flag_sn = True
    for i in range(4):
        transmit_inf[0] = 0x40
        transmit_inf[1] = 0xf7
        transmit_inf[2] = 0x5f
        transmit_inf[3] = 0x00 + i + 1
        bool_sn, transmit_sn = CAN_option.transmitCAN(0x602, transmit_inf, 0)
        if not bool_sn:
            return False

        while True:
            bool_sn, m_can_obj = CAN_option.receiveCANbyID(0x582, 2, 0)
            if bool_sn:
                break
            if bool_sn == 'stopReceive':
                return 'stopReceive', m_can_obj.Data
            elif not bool_sn:
                return False, m_can_obj.Data

        for j in range(4):
            if m_can_obj.Data[j + 4] != 0x00 and m_can_obj.Data[j + 4] != 0xff:
                sn_list.append(int(m_can_obj.Data[j + 4]))
            else:
                flag_sn = False
                break

        if not flag_sn:
            break

    # print('sn_list = ', sn_list)
    snCode = ''
    for asc in sn_list:
        snCode += chr(asc)
    # print('snCode = ', snCode)

    # 读设备rev
    rev_list = []
    flag_rev = True
    for i in range(4):
        transmit_inf[0] = 0x40
        transmit_inf[1] = 0xf9
        transmit_inf[2] = 0x5f
        transmit_inf[3] = 0x00 + i + 1
        bool_rev, transmit_rev = CAN_option.transmitCAN(0x602, transmit_inf, 0)
        if not bool_rev:
            return False

        while True:
            bool_rev, m_can_obj = CAN_option.receiveCANbyID(0x582, 2, 0)
            if bool_rev:
                break
            if bool_rev == 'stopReceive':
                return 'stopReceive', m_can_obj.Data
            elif not bool_rev:
                return False, m_can_obj.Data

        for j in range(4):
            if m_can_obj.Data[j + 4] != 0x00:
                rev_list.append(int(m_can_obj.Data[j + 4]))
            else:
                flag_rev = False
                break

        if not flag_rev:
            break

    # print('rev_list = ', rev_list)
    revCode = ''
    for asc in rev_list:
        revCode += chr(asc)
    # print('revCode = ', revCode)

    return True, [pnCode, snCode, revCode]

